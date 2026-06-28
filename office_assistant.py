# -*- coding: utf-8 -*-
"""万能办公助手 v6.1 商业版 — 模块化架构"""
import sys, os, warnings
warnings.filterwarnings("ignore")

if not getattr(sys, 'frozen', False):
    sys.path = [p for p in sys.path if 'hermes-agent' not in p and 'venv' not in p]

import re, json, time, threading, shutil, hashlib
from datetime import datetime, timedelta
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from collections import Counter
import base64, urllib.parse

# 导入工具模块
from utils import (
    DATA_DIR, CONFIG_FILE,
    PIL_AVAILABLE, OPENPYL_AVAILABLE, DOCX_AVAILABLE,
    PDF_AVAILABLE, REPORTLAB_AVAILABLE, PYPDFIUM_AVAILABLE,
    _PIL_ERROR, _OPENPYL_ERROR, _DOCX_ERROR,
    _PDF_ERROR, _REPORTLAB_ERROR, _PYPDFIUM_ERROR,
    load_config, save_config, safe_str, get_font,
    _log_error, safe_cond_check, get_lib_status_text,
    ProgressDialog, show_result, write_diagnostic_log
)
from lib_license import LicenseManager

# 写启动诊断
write_diagnostic_log()


class OfficeAssistant:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("万能办公助手 v6.1 商业版")
        self.root.geometry("1320x860")
        self.root.minsize(1050, 720)
        self.colors = {
            'primary': '#4F46E5', 'primary_hover': '#4338CA',
            'secondary': '#7C3AED', 'accent': '#06B6D4',
            'success': '#059669', 'warning': '#D97706',
            'danger': '#DC2626', 'dark': '#0F172A',
            'light': '#F8FAFC', 'white': '#FFFFFF',
            'gray': '#64748B', 'gray_light': '#94A3B8',
            'border': '#E2E8F0', 'hover_bg': '#F1F5F9',
            'nav_bg': '#0F172A', 'nav_item': '#334155',
            'nav_active': '#4F46E5', 'nav_text': '#94A3B8',
            'nav_text_active': '#FFFFFF', 'card_bg': '#FFFFFF',
            'success_bg': '#ECFDF5', 'warning_bg': '#FFFBEB',
        }
        self.clipboard_history = []
        self.clipboard_listbox = None
        self.memo_text = None
        self.todo_listbox = None
        self._todo_data = []
        self.config = load_config()
        self.progress_var = tk.StringVar(value="就绪")
        self.license_label = None
        self.license = LicenseManager(self)
        self._setup_ui()
        # 首次启动 → 欢迎；后续 → 许可证检查
        if not self.license.LICENSE_FILE.exists():
            self.root.after(500, self.license.show_welcome_dialog)
        else:
            self.root.after(1000, self.license.check_and_show_license_warning)

    def _setup_ui(self):
        self.root.configure(bg=self.colors['light'])
        # ==== 顶部标题栏 ====
        top_bar = tk.Frame(self.root, bg=self.colors['dark'], height=56)
        top_bar.pack(fill=tk.X, side=tk.TOP)
        top_bar.pack_propagate(False)
        tk.Label(top_bar, text="🧰  万能办公助手 v6.1", font=("微软雅黑", 15, "bold"),
                fg="white", bg=self.colors['dark']).pack(side=tk.LEFT, padx=20, pady=10)
        tk.Label(top_bar, text="商业版 · 万能办公助手 · 安全 · 高效 · 全能",
                font=("微软雅黑", 9), fg="#94A3B8", bg=self.colors['dark']).pack(side=tk.LEFT, padx=10, pady=10)
        tk.Button(top_bar, text="ℹ️ 关于", command=self.license.about_dialog, cursor="hand2",
                 font=("微软雅黑", 9), bd=0, padx=8, pady=2,
                 bg=self.colors['dark'], fg="#94A3B8",
                 activebackground='#1E293B', activeforeground='white').pack(side=tk.RIGHT, padx=(0, 12))
        # ==== 左侧导航 ====
        nav = tk.Frame(self.root, bg=self.colors['nav_bg'], width=220)
        nav.pack(side=tk.LEFT, fill=tk.Y)
        nav.pack_propagate(False)

        self.nav_buttons = {}
        nav_items = [
            ("📁  文件处理", "file"),
            ("📊  Excel处理", "excel"),
            ("📝  文档处理", "doc"),
            ("📄  PDF工具", "pdf"),
            ("🖼  图片工具", "image"),
            ("🔄  格式互转", "convert"),
            ("⚡  快捷工具", "quick"),
            ("📋  剪贴板", "clipboard"),
            ("🔍  内容搜索", "search"),
            ("📅  日程管理", "calendar"),
        ]
        for text, key in nav_items:
            btn = tk.Button(nav, text=text, font=("微软雅黑", 11), anchor="w", padx=24, pady=13,
                           bg=self.colors['nav_bg'], fg="#CBD5E1", bd=0, cursor="hand2",
                           activebackground=self.colors['nav_active'], activeforeground="white",
                           relief="flat",
                           command=lambda k=key: self._navigate(k))
            btn.pack(fill=tk.X)
            self.nav_buttons[key] = btn
        # ==== 主内容 ====
        main_bg = tk.Frame(self.root, bg=self.colors['light'])
        main_bg.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.content_frame = tk.Frame(main_bg, bg=self.colors['light'])
        self.content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        # ==== 底部状态栏 ====
        bottom = tk.Frame(self.root, bg=self.colors['dark'], height=32)
        bottom.pack(side=tk.BOTTOM, fill=tk.X)
        bottom.pack_propagate(False)
        tk.Label(bottom, textvariable=self.progress_var, fg="#CBD5E1", bg=self.colors['dark'],
                font=("Consolas", 9), anchor="w", padx=15).pack(side=tk.LEFT)
        # 右侧显示可用库
        libs = []
        if PIL_AVAILABLE: libs.append("PIL")
        if OPENPYL_AVAILABLE: libs.append("Openpyxl")
        if DOCX_AVAILABLE: libs.append("Docx")
        if PDF_AVAILABLE: libs.append("PDF")
        lib_text = "+".join(libs)
        tk.Label(bottom, text=lib_text, fg="#475569", bg=self.colors['dark'],
                font=("Consolas",8), padx=15).pack(side=tk.RIGHT)
        # 许可证状态标签（在状态栏中间）
        self.license_label = tk.Label(bottom, text="", fg="#10B981", bg=self.colors['dark'],
                font=("Consolas", 9), padx=10)
        self.license_label.pack(side=tk.RIGHT)
        self._navigate("file")
        self.root.after(500, self._update_status_bar_license)
        self.root.after(3000, self._auto_clip_timer)


    def _navigate(self, key):
        # 高亮当前
        for k, btn in self.nav_buttons.items():
            if k == key:
                btn.config(bg=self.colors['nav_active'], fg="white")
            else:
                btn.config(bg=self.colors['nav_bg'], fg="#CBD5E1")
        # 路由
        routes = {
            "file": self._show_file_tools,
            "excel": self._show_excel_tools,
            "doc": self._show_doc_tools,
            "pdf": self._show_pdf_tools,
            "image": self._show_image_tools,
            "convert": self._show_convert_tools,
            "quick": self._show_quick_tools,
            "clipboard": self._show_clipboard_tools,
            "search": self._show_search_tools,
            "calendar": self._show_calendar_tools,
        }
        routes.get(key, self._show_file_tools)()

    def set_status(self, msg):
        self.progress_var.set(msg)
        self.root.update_idletasks()

    def clear_content(self):
        for w in self.content_frame.winfo_children():
            w.destroy()

    def _run_thread(self, func, *args, done_msg="完成"):
        """后台线程执行（修复：之前是stub，线程代码被 _run_with_progress 隔断成了孤儿代码）"""
        self.set_status("处理中…")
        def worker():
            try:
                func(*args)
            except Exception as e:
                import traceback
                err_text = traceback.format_exc()
                try:
                    log_path = Path(DATA_DIR) / "crash.log"
                    with open(str(log_path), "a", encoding="utf-8") as f:
                        from datetime import datetime
                        f.write(f"\n{'='*50}\n[{datetime.now()}] _run_thread 异常\n{err_text}\n")
                except: pass
                self.root.after(0, lambda: messagebox.showerror("运行错误", f"错误详情:\n\n{err_text[-800:]}"))
            self.root.after(0, lambda: self.set_status(done_msg))
        threading.Thread(target=worker, daemon=True).start()

    def _run_with_progress(self, title, worker_func):
        """带进度对话框的后台线程执行"""
        pd = ProgressDialog(self.root, title=title)
        def wrapper():
            try:
                worker_func(pd)
            except Exception as e:
                import traceback
                err = traceback.format_exc()
                try:
                    log_path = Path(DATA_DIR) / "crash.log"
                    with open(str(log_path), "a", encoding="utf-8") as f:
                        f.write(f"\n[{datetime.now()}] {title} 异常\n{err}\n")
                except: pass
                pd.update(f"[ERROR] {e}", progress=1)
                def _show_err():
                    pd.close()
                    messagebox.showerror("错误", f"{title} 失败\n\n{str(e)[:300]}")
                self.root.after(500, _show_err)
        threading.Thread(target=wrapper, daemon=True).start()

    def _create_card(self, parent, title, desc, callback, btn_label="执行"):
        """统一按钮卡片（增强版：阴影边框+悬停动效）"""
        outer = tk.Frame(parent, bg=self.colors['border'], highlightthickness=0)
        outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 5), pady=(5, 5))
        card = tk.Frame(outer, bg='white', highlightthickness=0, relief='flat')
        card.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        tk.Frame(card, bg=self.colors['primary'], height=3).pack(fill=tk.X)
        body = tk.Frame(card, bg='white')
        body.pack(fill=tk.BOTH, expand=True, padx=14, pady=(16, 12))
        tk.Label(body, text=title, font=("微软雅黑", 11, "bold"), bg='white',
                fg=self.colors['dark']).pack(anchor='w')
        desc_lbl = tk.Label(body, text=desc, wraplength=170, font=("微软雅黑", 9),
                fg=self.colors['gray'], bg='white', justify=tk.LEFT)
        desc_lbl.pack(anchor='w', pady=(6, 12))
        btn = tk.Button(body, text=btn_label, command=callback, cursor="hand2",
                 bg=self.colors['primary'], fg="white", bd=0,
                 font=("微软雅黑", 9, "bold"), padx=16, pady=4,
                 activebackground=self.colors['primary_hover'])
        btn.pack(anchor='w')
        def _on_enter(e):
            outer.config(bg=self.colors['primary'])
            card.config(bg='#F8FAFF')
            for c in card.winfo_children():
                if not isinstance(c, tk.Button):
                    try: c.config(bg='#F8FAFF')
                    except: pass
            for c in body.winfo_children():
                if not isinstance(c, tk.Button):
                    try: c.config(bg='#F8FAFF')
                    except: pass
        def _on_leave(e):
            outer.config(bg=self.colors['border'])
            card.config(bg='white')
            for c in card.winfo_children():
                if not isinstance(c, tk.Button):
                    try: c.config(bg='white')
                    except: pass
            for c in body.winfo_children():
                if not isinstance(c, tk.Button):
                    try: c.config(bg='white')
                    except: pass
        for widget in [outer, card, body, desc_lbl]:
            widget.bind('<Enter>', _on_enter)
            widget.bind('<Leave>', _on_leave)
        def _click(e): callback()
        for widget in [outer, card, body]:
            widget.bind('<Button-1>', _click)
            for c in body.winfo_children():
                try: c.bind('<Button-1>', _click)
                except: pass


    def _show_file_tools(self):
        self.clear_content()
        self._section_header("文件处理工具", "批量操作 · 智能分类 · 查重去重")
        row = tk.Frame(self.content_frame, bg=self.colors['light'])
        row.pack(fill=tk.X, padx=10)
        self._create_card(row, " 📝 批量重命名", "前缀/后缀/替换/序号", self._batch_rename_dialog, "开始")
        self._create_card(row, " 📂 文件分类", "递归子目录整理", self._classify_files_dlg, "开始")
        self._create_card(row, " 🔎 查重并清理", "MD5查重+智能删除", self._find_dupes_dlg, "开始")

    def _batch_rename_dialog(self):
        files = filedialog.askopenfilenames(title="选择文件")
        if not files: return
        win = tk.Toplevel(self.root)
        win.title("批量重命名"); win.geometry("600x560"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"已选 {len(files)} 个文件", font=("微软雅黑", 11, "bold"),
                bg=self.colors['light']).pack(pady=10)
        # 模式
        mode_f = tk.LabelFrame(win, text="模式", font=("微软雅黑", 10), bg=self.colors['light'], padx=10, pady=5)
        mode_f.pack(fill=tk.X, padx=20, pady=5)
        mode_var = tk.StringVar(value="prefix")
        modes = [("prefix","前缀"),("suffix","后缀"),("replace","替换"),("number","序号")]
        for v,t in modes:
            tk.Radiobutton(mode_f, text=t, variable=mode_var, value=v, font=("微软雅黑", 10),
                          bg=self.colors['light']).pack(side=tk.LEFT, padx=6)
        # 输入
        inp = tk.Frame(win, bg=self.colors['light']); inp.pack(fill=tk.X, padx=20, pady=5)
        find_entry = tk.Entry(inp, font=("微软雅黑", 10)); find_entry.grid(row=0, column=1, sticky="ew")
        repl_entry = tk.Entry(inp, font=("微软雅黑", 10)); repl_entry.grid(row=1, column=1, sticky="ew")
        text_entry = tk.Entry(inp, font=("微软雅黑", 10)); text_entry.grid(row=0, column=1, sticky="ew")
        fl = tk.Label(inp, text="查找:", font=("微软雅黑",10), bg=self.colors['light'])
        rl = tk.Label(inp, text="替换:", font=("微软雅黑",10), bg=self.colors['light'])
        tl = tk.Label(inp, text="文字:", font=("微软雅黑",10), bg=self.colors['light'])
        def _show_inputs(*a):
            for w in inp.winfo_children(): w.grid_forget()
            m = mode_var.get()
            if m == "replace":
                fl.grid(row=0,column=0,sticky="w",padx=5); find_entry.grid(row=0,column=1,sticky="ew",padx=5)
                rl.grid(row=1,column=0,sticky="w",padx=5); repl_entry.grid(row=1,column=1,sticky="ew",padx=5)
                inp.columnconfigure(1,weight=1)
            elif m in ("prefix","suffix"):
                tl.grid(row=0,column=0,sticky="w",padx=5); text_entry.grid(row=0,column=1,sticky="ew",padx=5)
                inp.columnconfigure(1,weight=1)
            else:
                tk.Label(inp, text="自动添加 001,002…", font=("微软雅黑",10), fg=self.colors['gray'],
                        bg=self.colors['light']).grid(row=0,column=0,columnspan=2)
        mode_var.trace_add("write", _show_inputs); _show_inputs()
        log = scrolledtext.ScrolledText(win, height=14, font=("Consolas",9)); log.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)
        def _go():
            m = mode_var.get(); ok = skip = 0; log.delete("1.0", tk.END)
            try:
                if m == "prefix":
                    txt = text_entry.get()
                    if not txt: return messagebox.showwarning("提示", "请填写文字")
                    for fp in files:
                        p = Path(fp); new = p.parent / (txt + p.name)
                        if new.exists(): skip += 1; log.insert(tk.END, f"[SKIP] {new.name}\n")
                        else: p.rename(new); ok += 1; log.insert(tk.END, f"[OK] {p.name} → {new.name}\n")
                elif m == "suffix":
                    txt = text_entry.get()
                    if not txt: return messagebox.showwarning("提示", "请填写文字")
                    for fp in files:
                        p = Path(fp); s,e = p.stem, p.suffix
                        new = p.parent / f"{s}{txt}{e}"
                        if new.exists(): skip += 1
                        else: p.rename(new); ok += 1; log.insert(tk.END, f"[OK] {p.name} → {new.name}\n")
                elif m == "replace":
                    ft = find_entry.get(); rt = repl_entry.get()
                    if not ft: return messagebox.showwarning("提示", "请填写查找内容")
                    for fp in files:
                        p = Path(fp); nn = p.name.replace(ft, rt)
                        if nn == p.name: continue
                        new = p.parent / nn
                        if new.exists(): skip += 1
                        else: p.rename(new); ok += 1; log.insert(tk.END, f"[OK] {p.name} → {new.name}\n")
                else:
                    for i, fp in enumerate(files, 1):
                        p = Path(fp); s,e = p.stem, p.suffix
                        nn = f"{s}_{i:03d}{e}"; new = p.parent / nn
                        if new.exists(): skip += 1; continue
                        p.rename(new); ok += 1; log.insert(tk.END, f"[OK] {p.name} → {nn}\n")
            except Exception as e: log.insert(tk.END, f"错误: {e}\n")
            log.insert(tk.END, f"\n✅ {ok} 成功  |  ⏭ {skip} 跳过\n")
        tk.Button(win, text="▶ 开始重命名", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"), width=16).pack(pady=10)
        tk.Button(win, text="关闭", command=win.destroy, font=("微软雅黑", 9)).pack(pady=5)

    def _classify_files_dlg(self):
        d = filedialog.askdirectory(title="选择文件夹"); 
        if not d: return
        win = tk.Toplevel(self.root); win.title("分类整理"); win.geometry("550x450"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        recursive = tk.BooleanVar(value=True)
        tk.Checkbutton(win, text="递归子目录", variable=recursive, font=("微软雅黑", 10),
                      bg=self.colors['light']).pack(pady=10)
        log = scrolledtext.ScrolledText(win, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _go():
            cats = {"Documents":".doc|.docx|.pdf|.xls|.xlsx|.ppt|.pptx|.txt|.md|.csv|.rtf",
                    "Images":".jpg|.jpeg|.png|.gif|.bmp|.svg|.webp|.ico|.tiff",
                    "Videos":".mp4|.avi|.mkv|.mov|.wmv|.flv|.webm",
                    "Audio":".mp3|.wav|.flac|.aac|.ogg|.wma|.m4a",
                    "Archives":".zip|.rar|.7z|.tar|.gz|.bz2",
                    "Code":".py|.js|.java|.cpp|.c|.cs|.go|.html|.css|.json|.yaml|.xml|.sh"}
            cat_map = {c: set(v.split("|")) for c,v in cats.items()}
            stats = {}; log.delete("1.0", tk.END)
            it = Path(d).rglob("*") if recursive.get() else Path(d).iterdir()
            for f in it:
                if not f.is_file(): continue
                # skip already in category dirs
                if f.parent.name in cats: continue
                ext = f.suffix.lower()
                tgt = next((c for c, exts in cat_map.items() if ext in exts), "Others")
                td = Path(d) / tgt; td.mkdir(exist_ok=True)
                try:
                    dest = td / f.name
                    if dest.exists():
                        dest = td / f"{f.stem}_{int(time.time())}{f.suffix}"
                    shutil.move(str(f), str(dest))
                    stats[tgt] = stats.get(tgt, 0) + 1
                    log.insert(tk.END, f"[OK] {f.name} → {tgt}/\n")
                except Exception as e: log.insert(tk.END, f"[FAIL] {f.name}: {e}\n")
            log.insert(tk.END, "\n=== 📊 统计 ===\n")
            for c,n in sorted(stats.items()): log.insert(tk.END, f"  {c}: {n}个\n")
            self.set_status("完成")

    def _show_success_dialog(self, title, message, file_path=None):
        """成功提示对话框（带「打开文件夹」按钮）"""
        win = tk.Toplevel(self.root)
        win.title(title)
        win.geometry("440x200")
        win.transient(self.root)
        win.grab_set()
        win.configure(bg='#FFFFFF')
        tk.Frame(win, bg='#4F46E5', height=3).pack(fill=tk.X)
        tk.Label(win, text="✅ " + title, font=("微软雅黑", 14, "bold"),
                 fg='#059669', bg='#FFFFFF').pack(pady=(15, 8))
        msg = tk.Label(win, text=message, font=("微软雅黑", 10),
                 fg='#374151', bg='#FFFFFF', wraplength=400, justify=tk.LEFT)
        msg.pack(pady=5, padx=20)
        btn_f = tk.Frame(win, bg='#FFFFFF')
        btn_f.pack(pady=(12, 18))
        if file_path:
            import os as _os
            tk.Button(btn_f, text="📂 打开文件夹", command=lambda p=file_path: _os.startfile(str(Path(p).parent)),
                     cursor="hand2", bg='#E2E8F0', fg='#0F172A',
                     font=("微软雅黑", 10), width=14, bd=0).pack(side=tk.LEFT, padx=6)
        tk.Button(btn_f, text="确定", command=win.destroy, cursor="hand2",
                 bg='#4F46E5', fg='white', font=("微软雅黑", 10),
                 width=10, bd=0).pack(side=tk.LEFT, padx=6)

        tk.Button(win, text="▶ 开始分类", command=lambda: self._run_thread(_go, done_msg="分类完成"),
                 cursor="hand2", bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"), width=18).pack(pady=10)

    def _find_dupes_dlg(self):
        d = filedialog.askdirectory(title="选择目录"); 
        if not d: return
        win = tk.Toplevel(self.root); win.title("重複文件清理"); win.geometry("650x470"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        log = scrolledtext.ScrolledText(win, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        dup_map = {}  # hash: list of paths
        def _go():
            log.delete("1.0", tk.END); log.insert(tk.END, "🔍 扫描中…\n"); self.set_status("扫描文件…"); win.update()
            hash_dict = {}; count = 0
            for f in Path(d).rglob("*"):
                if not f.is_file(): continue
                count += 1
                if count % 30 == 0: win.update()
                try:
                    size = f.stat().st_size
                    # skip large files (> 200MB)
                    if size > 200 * 1024 * 1024: continue
                    with open(f, 'rb') as fh:
                        h = hashlib.md5(fh.read()).hexdigest()
                    if h in hash_dict:
                        if h not in dup_map: dup_map[h] = {hash_dict[h]}
                        dup_map[h].add(f)
                    else: hash_dict[h] = f
                except Exception: pass
            # show results
            total_dup = sum(len(v) for v in dup_map.values())
            log.insert(tk.END, f"📊 扫描 {count} 个文件 → 发现 {len(dup_map)} 组重复 ({total_dup} 个文件)\n\n")
            for h, paths in dup_map.items():
                sizes = [p.stat().st_size for p in paths]
                log.insert(tk.END, f"▸ {len(paths)}份重复\n")
                for p in sorted(paths, key=lambda x: x.stat().st_mtime, reverse=True):
                    log.insert(tk.END, f"   📄 {p}  ({p.stat().st_size // 1024}KB)\n")
                log.insert(tk.END, "\n")
            log.insert(tk.END, "💡 按「删除旧版」将保留每个重复集中最新修改的文件\n")
            self.set_status("扫描完成")
        tk.Button(win, text="🔍 开始扫描", command=lambda: self._run_thread(_go, done_msg="扫描完成"),
                 cursor="hand2", bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"), width=14).pack(pady=8)
        def _delete_old():
            if not dup_map: return messagebox.showinfo("提示", "请先扫描")
            if not messagebox.askyesno("确认", f"将删除 {len(dup_map)} 组重复中的旧文件，保留最新修改版\n\n确定？"): return
            log.insert(tk.END, "🗑 删除旧版…\n"); deleted = 0
            for h, paths in dup_map.items():
                sorted_paths = sorted(paths, key=lambda x: x.stat().st_mtime, reverse=True)
                keep = sorted_paths[0]
                for p in sorted_paths[1:]:
                    try:
                        p.unlink()
                        deleted += 1
                        log.insert(tk.END, f"[DEL] {p}\n")
                    except Exception as e: log.insert(tk.END, f"[FAIL] {p}: {e}\n")
            log.insert(tk.END, f"\n✅ 删除 {deleted} 个重复文件\n")
        tk.Button(win, text="🗑 删除旧版（保留最新）", command=_delete_old, cursor="hand2",
                 bg=self.colors['danger'], fg="white", font=("微软雅黑", 10, "bold"), width=24).pack(pady=5)

    def _section_header(self, title, sub=""):
        f = tk.Frame(self.content_frame, bg=self.colors['light']); f.pack(fill=tk.X, padx=20, pady=(15,5))
        tk.Label(f, text=title, font=("微软雅黑", 18, "bold"), bg=self.colors['light'], fg=self.colors['dark']).pack(anchor="w")
        if sub: tk.Label(f, text=sub, font=("微软雅黑", 10), bg=self.colors['light'], fg=self.colors['gray']).pack(anchor="w")


    # ==================== Excel处理 ====================
    def _show_excel_tools(self):
        self.clear_content()
        self._section_header("Excel处理工具", "合并 · 拆分 · 筛选 · 统计 · 导出")
        if not OPENPYL_AVAILABLE:
            tk.Label(self.content_frame, text=f"⚠️ openpyxl未安装\n{_OPENPYL_ERROR}", fg="red",
                    font=("微软雅黑", 12), bg=self.colors['light']).pack(pady=60)
            return
        for rdata in [
            [("Excel合并","多文件→一个文件", self._merge_excel_dlg),
             ("Excel拆分","按行数/按Sheet", self._split_excel_dlg),
             ("按条件筛选","保存为新文件", self._filter_excel_dlg)],
            [("数据统计","列级统计+频次分布", self._stats_excel_dlg),
             ("导出为CSV","UTF-8-BOM", self._excel_to_csv_dlg)],
        ]:
            row = tk.Frame(self.content_frame, bg=self.colors['light'])
            row.pack(fill=tk.X, padx=10)
            for title, desc, cb in rdata:
                self._create_card(row, title, desc, cb)

    def _merge_excel_dlg(self):
        files = filedialog.askopenfilenames(title="选择Excel文件", filetypes=[("Excel","*.xlsx")])
        if not files: return
        save = filedialog.asksaveasfilename(title="保存合并结果", defaultextension=".xlsx", filetypes=[("Excel","*.xlsx")])
        if not save: return
        def _go():
            wb_out = openpyxl.Workbook()
            ws_out = wb_out.active; ws_out.title = "合并"
            first = True; total = 0
            for fp in files:
                wb = openpyxl.load_workbook(fp, read_only=True)
                for ws in wb.worksheets:
                    rows = list(ws.iter_rows(values_only=True))
                    if not rows: continue
                    if first:
                        for r in rows: ws_out.append(r); total += 1
                        first = False
                    else:
                        for r in rows[1:]: ws_out.append(r); total += 1
                wb.close()
            wb_out.save(save); wb_out.close()
            self.root.after(0, lambda: self._show_success_dialog("完成", f"合并 {len(files)} 文件\n总计 {total} 行"))
        self._run_thread(_go, done_msg="合并完成")

    def _split_excel_dlg(self):
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel","*.xlsx")])
        if not fp: return
        out_dir = filedialog.askdirectory(title="输出目录")
        if not out_dir: return
        win = tk.Toplevel(self.root); win.title("拆分"); win.geometry("380x180"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="每文件行数:", font=("微软雅黑", 11), bg=self.colors['light']).pack(pady=15)
        e = tk.Entry(win, font=("微软雅黑", 12), width=10); e.insert(0, "500"); e.pack()
        def _go():
            try: rows_per = int(e.get()); assert rows_per > 0
            except: messagebox.showwarning("提示","请输入正整数"); return
            win.destroy()
            def _work():
                src = openpyxl.load_workbook(fp, read_only=True)
                all_rows = list(src.active.iter_rows(values_only=True)); src.close()
                if len(all_rows)<2: return messagebox.showinfo("提示","文件无数据")
                hdr = all_rows[0]; data = all_rows[1:]; stem = Path(fp).stem
                for i in range(0, len(data), rows_per):
                    chunk = data[i:i+rows_per]
                    wb2 = openpyxl.Workbook(); ws2 = wb2.active
                    ws2.append(hdr)
                    for r in chunk: ws2.append(r)
                    op = Path(out_dir) / f"{stem}_第{i//rows_per+1}组.xlsx"
                    wb2.save(op); wb2.close()
                self.root.after(0, lambda: self._show_success_dialog("完成", f"拆分完成 → {out_dir}"))
            self._run_thread(_work, done_msg="拆分完成")
        tk.Button(win, text="▶ 开始拆分", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=14).pack(pady=10)

    def _filter_excel_dlg(self):
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel","*.xlsx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存筛选结果", defaultextension=".xlsx", filetypes=[("Excel","*.xlsx")])
        if not save: return
        try:
            wb = openpyxl.load_workbook(fp, read_only=True)
            all_rows = list(wb.active.iter_rows(values_only=True)); wb.close()
            if len(all_rows)<2: return messagebox.showwarning("提示","无数据")
            hdr = all_rows[0]; data = all_rows[1:]
            win = tk.Toplevel(self.root); win.title("筛选"); win.geometry("550x520"); win.transient(self.root); win.grab_set()
            win.configure(bg=self.colors['light'])
            tk.Label(win, text="选择列:", font=("微软雅黑", 10), bg=self.colors['light']).pack(pady=5)
            lb = tk.Listbox(win, height=10, font=("微软雅黑", 10))
            for i,h in enumerate(hdr): lb.insert(tk.END, f"[{i}] {h}")
            lb.pack(fill=tk.X, padx=20)
            tk.Label(win, text="筛选条件 (包含匹配):", font=("微软雅黑", 10), bg=self.colors['light']).pack(pady=5)
            ve = tk.Entry(win, width=40, font=("微软雅黑", 10)); ve.pack()
            log = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
            def _go():
                sel = lb.curselection()
                if not sel: return messagebox.showwarning("提示","请选择列")
                ci = sel[0]; v = ve.get().strip()
                if not v: return messagebox.showwarning("提示","请填写条件")
                fltd = [hdr]
                for row in data:
                    if ci < len(row) and row[ci] is not None and v.lower() in safe_str(row[ci]).lower():
                        fltd.append(row)
                wb2 = openpyxl.Workbook(); ws2 = wb2.active; ws2.title = "筛选结果"
                for r in fltd: ws2.append(r)
                wb2.save(save); wb2.close()
                log.delete("1.0", tk.END); log.insert(tk.END, f"列: {hdr[ci]}  条件: {v}\n匹配: {len(fltd)-1}行\n\n保存: {save}")
            tk.Button(win, text="▶ 执行筛选", command=_go, cursor="hand2",
                     bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=14).pack(pady=8)
        except Exception as e: messagebox.showerror("错误", str(e))

    def _stats_excel_dlg(self):
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel","*.xlsx")])
        if not fp: return
        try:
            wb = openpyxl.load_workbook(fp, read_only=True)
            all_rows = list(wb.active.iter_rows(values_only=True)); wb.close()
            if len(all_rows)<2: return messagebox.showwarning("提示","无数据")
            hdr = all_rows[0]; data = all_rows[1:]
            win = tk.Toplevel(self.root); win.title("统计"); win.geometry("750x520"); win.transient(self.root)
            win.configure(bg=self.colors['light'])
            log = scrolledtext.ScrolledText(win, font=("Consolas", 10)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
            log.insert(tk.END, "="*50 + "\n"
                       f"📄 {Path(fp).name}\n"
                       f"   行数: {len(all_rows)} (含表头) | 数据行: {len(data)} | 列数: {len(hdr)}\n"
                       "="*50 + "\n\n")
            for ci, h in enumerate(hdr):
                vals = [row[ci] for row in data if ci < len(row) and row[ci] is not None]
                if not vals: log.insert(tk.END, f"📌 [{h}] 空列\n"); continue
                nums = []
                for v in vals:
                    try: nums.append(float(v))
                    except (ValueError,TypeError): pass
                if nums:
                    mn,mx = min(nums),max(nums)
                    sd = (sum((n-sum(nums)/len(nums))**2 for n in nums)/len(nums))**0.5 if len(nums)>1 else 0
                    log.insert(tk.END, f"📊 [{h}] 数值列 ({len(nums)}条)\n"
                                f"   总和={sum(nums):.2f}  平均={sum(nums)/len(nums):.2f}  标准差={sd:.2f}\n"
                                f"   最大={mx:.2f}  最小={mn:.2f}  极差={mx-mn:.2f}\n\n")
                else:
                    c = Counter(safe_str(v) for v in vals)
                    log.insert(tk.END, f"📝 [{h}] 文本列 ({len(vals)}条)  唯一值: {len(c)}\n")
                    for v,n in c.most_common(5):
                        pct = n/len(vals)*100
                        log.insert(tk.END, f"   {v}: {n}次 ({pct:.1f}%)\n")
                    log.insert(tk.END, "\n")
        except Exception as e: messagebox.showerror("错误", str(e))

    def _excel_to_csv_dlg(self):
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel","*.xlsx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存CSV", defaultextension=".csv", filetypes=[("CSV","*.csv")])
        if not save: return
        def _go():
            import csv
            wb = openpyxl.load_workbook(fp, read_only=True)
            with open(save, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                for ws in wb.worksheets:
                    writer.writerow([f"=== Sheet: {ws.title} ==="])
                    for row in ws.iter_rows(values_only=True):
                        writer.writerow([safe_str(c) for c in row])
            wb.close()
            self.root.after(0, lambda: self._show_success_dialog("完成", f"已导出\n{save}"))
        self._run_thread(_go, done_msg="导出完成")


    # ==================== 文档处理 ====================
    def _show_doc_tools(self):
        self.clear_content()
        self._section_header("文档处理工具", "Word → 文本 · 表格提取 · 合并 · 对比")
        if not DOCX_AVAILABLE:
            tk.Label(self.content_frame, text="⚠️ python-docx未安装", fg="red",
                    font=("微软雅黑", 12), bg=self.colors['light']).pack(pady=60)
            return
        row = tk.Frame(self.content_frame, bg=self.colors['light']); row.pack(fill=tk.X, padx=10)
        self._create_card(row, "📄 Word→文本", "提取纯文本+表格", self._word_to_text_dlg)
        self._create_card(row, "📚 批量Word→文本", "多文件合并为文本", self._batch_word_dlg)
        self._create_card(row, "📊 表格→Excel", "提取Word表格到Excel", self._extract_tables_dlg)
        row2 = tk.Frame(self.content_frame, bg=self.colors['light']); row2.pack(fill=tk.X, padx=10)
        self._create_card(row2, "🔗 合并Word", "多文档合为一个", self._merge_word_dlg)
        self._create_card(row2, "🔍 对比Word", "两文档差异高亮", self._diff_word_dlg)

    def _word_to_text_dlg(self):
        fp = filedialog.askopenfilename(title="选择Word", filetypes=[("Word","*.docx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存文本", defaultextension=".txt", filetypes=[("Text","*.txt")])
        if not save: return
        def _go():
            doc = Document(fp); lines = []
            for p in doc.paragraphs:
                if p.text.strip(): lines.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    rt = "\t".join(cell.text for cell in row.cells)
                    if rt.strip(): lines.append(rt)
            Path(save).write_text("\n".join(lines), encoding="utf-8")
            self.root.after(0, lambda: self._show_success_dialog("完成", f"提取 {len(lines)} 行\n{save}"))
        self._run_thread(_go, done_msg="提取完成")

    def _batch_word_dlg(self):
        files = filedialog.askopenfilenames(title="选择Word文件", filetypes=[("Word","*.docx")])
        if not files: return
        save = filedialog.asksaveasfilename(title="保存合并", defaultextension=".txt", filetypes=[("Text","*.txt")])
        if not save: return
        def _go():
            out = []
            for fp in files:
                doc = Document(fp)
                out.append(f"\n========== {Path(fp).name} ==========\n")
                for p in doc.paragraphs:
                    if p.text.strip(): out.append(p.text)
            Path(save).write_text("\n".join(out), encoding="utf-8")
            self.root.after(0, lambda: self._show_success_dialog("完成", f"处理 {len(files)} 个文件\n{save}"))
        self._run_thread(_go, done_msg="批量提取完成")

    def _extract_tables_dlg(self):
        fp = filedialog.askopenfilename(title="选择Word", filetypes=[("Word","*.docx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存Excel", defaultextension=".xlsx", filetypes=[("Excel","*.xlsx")])
        if not save: return
        def _go():
            doc = Document(fp)
            if not doc.tables: self.root.after(0, lambda: messagebox.showinfo("提示","未发现表格")); return
            wb = openpyxl.Workbook(); ws = wb.active; ws.title = "提取结果"; ro = 0
            for ti, table in enumerate(doc.tables):
                ws.cell(row=ro+1, column=1, value=f"=== Table {ti+1} ===").font = Font(bold=True); ro += 1
                for row in table.rows:
                    for ci, cell in enumerate(row.cells):
                        c = ws.cell(row=ro+1, column=ci+1, value=cell.text)
                        c.alignment = Alignment(wrap_text=True)
                    ro += 1
                ro += 1
            wb.save(save); wb.close()
            self.root.after(0, lambda: self._show_success_dialog("完成", f"提取 {len(doc.tables)} 个表格\n{save}"))
        self._run_thread(_go, done_msg="提取表格完成")

    # ==================== PDF工具 ====================
    def _show_pdf_tools(self):
        self.clear_content()
        self._section_header("PDF工具", "提取文本 · 合并 · 拆分 · 加密 · 解密")
        if not PDF_AVAILABLE:
            err_msg = _PDF_ERROR
            tk.Label(self.content_frame, text=f"⚠️ PDF模块加载失败\n\n错误信息:\n{err_msg}\n\n请确保 PyPDF2 已正确安装",
                    fg="red", font=("微软雅黑", 12), bg=self.colors['light'], justify=tk.LEFT).pack(pady=60)
            return
        for rdata in [
            [("📄 PDF→文本","提取PDF纯文本", self._pdf_to_text_dlg),
             ("📑 合并PDF","多文件合并为一个", self._merge_pdf_dlg),
             ("✂ 拆分PDF","按页数拆分", self._split_pdf_dlg)],
            [("🔐 加密PDF","设置密码保护", self._encrypt_pdf_dlg),
             ("🔓 解密PDF","移除密码", self._decrypt_pdf_dlg)],
        ]:
            row = tk.Frame(self.content_frame, bg=self.colors['light']); row.pack(fill=tk.X, padx=10)
            for title, desc, cb in rdata: self._create_card(row, title, desc, cb)

    def _pdf_to_text_dlg(self):
        fp = filedialog.askopenfilename(title="选择PDF", filetypes=[("PDF","*.pdf")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存文本", defaultextension=".txt", filetypes=[("Text","*.txt")])
        if not save: return
        def _go():
            all_text = []
            with pdfplumber.open(fp) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t: all_text.append(f"--- Page {i+1} ---\n{t}")
            Path(save).write_text("\n\n".join(all_text), encoding="utf-8")
            self.root.after(0, lambda: self._show_success_dialog("完成", f"提取 {len(all_text)} 页\n{save}"))
        self._run_thread(_go, done_msg="PDF提取完成")

    def _merge_pdf_dlg(self):
        files = filedialog.askopenfilenames(title="选择PDF文件", filetypes=[("PDF","*.pdf")])
        if len(files)<2: return messagebox.showinfo("提示","至少选择2个PDF")
        save = filedialog.asksaveasfilename(title="保存合并PDF", defaultextension=".pdf", filetypes=[("PDF","*.pdf")])
        if not save: return
        def _go():
            merger = PyPDF2.PdfMerger()
            for fp in files: merger.append(fp)
            merger.write(save); merger.close()
            self.root.after(0, lambda: self._show_success_dialog("完成", f"合并 {len(files)} 个PDF\n{save}"))
        self._run_thread(_go, done_msg="PDF合并完成")

    def _split_pdf_dlg(self):
        fp = filedialog.askopenfilename(title="选择PDF", filetypes=[("PDF","*.pdf")])
        if not fp: return
        out_dir = filedialog.askdirectory(title="输出目录")
        if not out_dir: return
        win = tk.Toplevel(self.root); win.title("拆分PDF"); win.geometry("350x160"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="每份文件页数:", font=("微软雅黑", 11), bg=self.colors['light']).pack(pady=15)
        e = tk.Entry(win, font=("微软雅黑", 12), width=8); e.insert(0, "1"); e.pack()
        def _go():
            try: pp = int(e.get()); assert pp>0
            except: return messagebox.showwarning("提示","请输入正整数")
            win.destroy()
            def _work():
                reader = PyPDF2.PdfReader(fp); total = len(reader.pages); stem = Path(fp).stem
                for i in range(0, total, pp):
                    writer = PyPDF2.PdfWriter()
                    for j in range(i, min(i+pp, total)):
                        writer.add_page(reader.pages[j])
                    op = Path(out_dir) / f"{stem}_p{i+1}-{min(i+pp,total)}.pdf"
                    with open(op, 'wb') as f: writer.write(f)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"{total}页拆分为{ (total+pp-1)//pp }份"))
            self._run_thread(_work, done_msg="PDF拆分完成")
        tk.Button(win, text="▶ 开始拆分", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=14).pack(pady=10)

    def _encrypt_pdf_dlg(self):
        fp = filedialog.askopenfilename(title="选择PDF", filetypes=[("PDF","*.pdf")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存加密PDF", defaultextension=".pdf", filetypes=[("PDF","*.pdf")])
        if not save: return
        win = tk.Toplevel(self.root); win.title("设置密码"); win.geometry("350x180"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="密码:", font=("微软雅黑", 11), bg=self.colors['light']).pack(pady=15)
        e = tk.Entry(win, font=("微软雅黑", 12), show="*"); e.pack()
        def _go():
            pwd = e.get()
            if not pwd: return messagebox.showwarning("提示","请设置密码")
            win.destroy()
            def _work():
                reader = PyPDF2.PdfReader(fp); writer = PyPDF2.PdfWriter()
                for page in reader.pages: writer.add_page(page)
                writer.encrypt(pwd)
                with open(save, 'wb') as f: writer.write(f)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"已加密\n{save}"))
            self._run_thread(_work, done_msg="PDF加密完成")
        tk.Button(win, text="🔐 加密", command=_go, cursor="hand2",
                 bg=self.colors['danger'], fg="white", font=("微软雅黑", 11), width=12).pack(pady=10)

    def _decrypt_pdf_dlg(self):
        fp = filedialog.askopenfilename(title="选择PDF", filetypes=[("PDF","*.pdf")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存解密PDF", defaultextension=".pdf", filetypes=[("PDF","*.pdf")])
        if not save: return
        win = tk.Toplevel(self.root); win.title("输入密码"); win.geometry("350x180"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="密码:", font=("微软雅黑", 11), bg=self.colors['light']).pack(pady=15)
        e = tk.Entry(win, font=("微软雅黑", 12), show="*"); e.pack()
        def _go():
            pwd = e.get(); win.destroy()
            def _work():
                reader = PyPDF2.PdfReader(fp)
                if reader.is_encrypted:
                    reader.decrypt(pwd)
                writer = PyPDF2.PdfWriter()
                for page in reader.pages: writer.add_page(page)
                with open(save, 'wb') as f: writer.write(f)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"已解密\n{save}"))
            self._run_thread(_work, done_msg="PDF解密完成")
        tk.Button(win, text="🔓 解密", command=_go, cursor="hand2",
                 bg=self.colors['success'], fg="white", font=("微软雅黑", 11), width=12).pack(pady=10)


    # ================== 图片工具 ==================
    def _show_image_tools(self):
        self.clear_content()
        self._section_header("图片工具", "格式转换 · 缩放 · 水印 · 拼接 · 压缩 · OCR · 九宫格")
        if not PIL_AVAILABLE:
            tk.Label(self.content_frame, text="⚠️ Pillow未安装", fg="red",
                    font=("微软雅黑", 12), bg=self.colors['light']).pack(pady=60)
            return
        for rdata in [
            [("🖼 格式转换","PNG/JPG/BMP/WEBP", self._convert_img_dlg),
             ("📐 批量缩放","保持比例/自定义", self._resize_img_dlg),
             ("💧 添加水印","5位置/自定义文字", self._watermark_dlg)],
            [("🔗 图片拼接","垂直/水平", self._concat_img_dlg),
             ("📦 批量压缩","质量/尺寸压缩", self._compress_img_dlg),
             ("🔤 OCR识别","图片转文字", self._ocr_img_dlg)],
            [("📄 PDF转图","PDF页面转图片", self._pdf_to_img_dlg),
             ("📱 九宫格切图","1图切9格", self._nine_grid_dlg)],
        ]:
            row = tk.Frame(self.content_frame, bg=self.colors['light']); row.pack(fill=tk.X, padx=10)
            for title, desc, cb in rdata: self._create_card(row, title, desc, cb)

    def _convert_img_dlg(self):
        files = filedialog.askopenfilenames(title="选择图片",
            filetypes=[("图片","*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.webp")])
        if not files: return
        win = tk.Toplevel(self.root); win.title("格式转换"); win.geometry("420x280"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"已选 {len(files)} 张", font=("微软雅黑", 11, "bold"), bg=self.colors['light']).pack(pady=10)
        fmt_var = tk.StringVar(value="PNG")
        ttk.Combobox(win, textvariable=fmt_var, values=["PNG","JPEG","BMP","GIF","TIFF","WEBP"],
                    state="readonly", width=12).pack(pady=5)
        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _go():
            tf = fmt_var.get(); ok=fail=0; log.delete("1.0", tk.END)
            for fp in files:
                try:
                    img = Image.open(fp)
                    np = Path(fp).with_suffix(f".{tf.lower()}")
                    if img.mode=="RGBA" and tf=="JPEG": img = img.convert("RGB")
                    img.save(np); ok+=1; log.insert(tk.END, f"[OK] {Path(fp).name} → {np.name}\n")
                except Exception as e: fail+=1; log.insert(tk.END, f"[FAIL] {Path(fp).name}: {e}\n")
            log.insert(tk.END, f"\n✅ {ok}成功  ❌ {fail}失败\n")
        tk.Button(win, text="▶ 转换", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=12).pack(pady=8)

    def _resize_img_dlg(self):
        files = filedialog.askopenfilenames(title="选择图片",
            filetypes=[("图片","*.png *.jpg *.jpeg *.bmp *.gif")])
        if not files: return
        win = tk.Toplevel(self.root); win.title("批量缩放"); win.geometry("450x320"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack(pady=10)
        tk.Label(f1, text="W:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=0,col=0,padx=5)
        wv = tk.StringVar(value="800"); tk.Entry(f1, textvariable=wv, width=6, font=("微软雅黑", 10)).grid(row=0,col=1)
        tk.Label(f1, text="H:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=0,col=2,padx=5)
        hv = tk.StringVar(value="600"); tk.Entry(f1, textvariable=hv, width=6, font=("微软雅黑", 10)).grid(row=0,col=3)
        keep = tk.BooleanVar(value=True)
        tk.Checkbutton(win, text="保持比例", variable=keep, font=("微软雅黑", 10), bg=self.colors['light']).pack()
        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _go():
            try: w,h = int(wv.get()), int(hv.get()); assert w>0 and h>0
            except: return messagebox.showwarning("提示","请输入正整数")
            ok=fail=0; log.delete("1.0", tk.END)
            for fp in files:
                try:
                    img = Image.open(fp); ow,oh = img.size
                    if keep.get():
                        r = min(w/ow, h/oh); nw,nh = max(1,int(ow*r)), max(1,int(oh*r))
                    else: nw,nh = w,h
                    out = fp.parent / f"{fp.stem}_{nw}x{nh}{fp.suffix}"
                    img.resize((nw,nh), Image.LANCZOS).save(out)
                    ok+=1; log.insert(tk.END, f"[OK] {Path(fp).name} → {out.name}\n")
                except Exception as e: fail+=1; log.insert(tk.END, f"[FAIL] {Path(fp).name}: {e}\n")
            log.insert(tk.END, f"\n✅ {ok}成功\n")
        tk.Button(win, text="▶ 缩放", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=12).pack(pady=5)

    def _watermark_dlg(self):
        files = filedialog.askopenfilenames(title="选择图片",
            filetypes=[("图片","*.png *.jpg *.jpeg *.bmp")])
        if not files: return
        win = tk.Toplevel(self.root); win.title("添加水印"); win.geometry("520x420"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack(pady=10, padx=20, fill=tk.X)
        tk.Label(f1, text="水印文字:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=0,col=0,sticky="w")
        tv = tk.StringVar(value="万能办公助手"); tk.Entry(f1, textvariable=tv, width=30, font=("微软雅黑", 10)).grid(row=0,col=1,padx=5)
        tk.Label(f1, text="位置:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=1,col=0,sticky="w",pady=5)
        pv = tk.StringVar(value="右下"); ttk.Combobox(f1, textvariable=pv,
            values=["左上","右上","左下","右下","居中"], state="readonly", width=8).grid(row=1,col=1,sticky="w",padx=5)
        tk.Label(f1, text="字号:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=2,col=0,sticky="w",pady=5)
        sz = tk.StringVar(value="50"); tk.Entry(f1, textvariable=sz, width=6, font=("微软雅黑", 10)).grid(row=2,col=1,sticky="w",padx=5)
        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _go():
            try: size = int(sz.get()); assert size>0
            except: return messagebox.showwarning("提示","请输入有效字号")
            text = tv.get(); pos = pv.get()
            if not text: return messagebox.showwarning("提示","请输入水印文字")
            ok=fail=0; log.delete("1.0", tk.END); font = get_font(size)
            for fp in files:
                try:
                    img = Image.open(fp).convert("RGBA"); ov = Image.new("RGBA", img.size, (0,0,0,0))
                    d = ImageDraw.Draw(ov)
                    bbox = d.textbbox((0,0), text, font=font)
                    tw,th = bbox[2]-bbox[0], bbox[3]-bbox[1]; m=30
                    if pos=="左上": xy=(m,m)
                    elif pos=="右上": xy=(img.width-tw-m,m)
                    elif pos=="左下": xy=(m,img.height-th-m)
                    elif pos=="右下": xy=(img.width-tw-m,img.height-th-m)
                    else: xy=((img.width-tw)//2, (img.height-th)//2)
                    d.text(xy, text, fill=(255,255,255,220), font=font)
                    res = Image.alpha_composite(img, ov)
                    res.save(fp.parent / f"{fp.stem}_水印{fp.suffix}")
                    ok+=1; log.insert(tk.END, f"[OK] {Path(fp).name}\n")
                except Exception as e: fail+=1; log.insert(tk.END, f"[FAIL] {Path(fp).name}: {e}\n")
            log.insert(tk.END, f"\n✅ {ok}张完成\n")
        tk.Button(win, text="▶ 添加水印", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=14).pack(pady=8)

    def _concat_img_dlg(self):
        files = filedialog.askopenfilenames(title="选择图片(2+)",
            filetypes=[("图片","*.png *.jpg *.jpeg *.bmp")])
        if len(files)<2: return messagebox.showinfo("提示","至少选2张")
        save = filedialog.asksaveasfilename(title="保存拼接结果", defaultextension=".jpg",
            filetypes=[("JPEG","*.jpg"),("PNG","*.png")])
        if not save: return
        win = tk.Toplevel(self.root); win.title("拼接方向"); win.geometry("320x150"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        dv = tk.StringVar(value="vertical")
        f = tk.Frame(win, bg=self.colors['light']); f.pack(pady=20)
        tk.Label(f, text="方向:", font=("微软雅黑", 11), bg=self.colors['light']).pack(side=tk.LEFT)
        ttk.Combobox(f, textvariable=dv, values=["vertical","horizontal"], state="readonly", width=10).pack(side=tk.LEFT,padx=8)
        def _go():
            d = dv.get(); win.destroy()
            def _work():
                imgs = [Image.open(p).convert("RGB") for p in files]
                if d=="vertical":
                    mw = max(i.width for i in imgs); th = sum(i.height for i in imgs)
                    r = Image.new("RGB",(mw,th),(255,255,255)); y=0
                    for img in imgs: r.paste(img,((mw-img.width)//2,y)); y+=img.height
                else:
                    mh = max(i.height for i in imgs); tw = sum(i.width for i in imgs)
                    r = Image.new("RGB",(tw,mh),(255,255,255)); x=0
                    for img in imgs: r.paste(img,(x,(mh-img.height)//2)); x+=img.width
                r.save(save)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"拼接 {len(files)} 张\n{save}"))
            self._run_thread(_work, done_msg="拼接完成")
        tk.Button(win, text="▶ 开始拼接", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11), width=14).pack(pady=10)


    # ==================== 快捷工具 ====================
    def _show_quick_tools(self):
        self.clear_content()
        self._section_header("快捷工具", "文本 · 编码 · 时间 · 哈希")
        row = tk.Frame(self.content_frame, bg=self.colors['light']); row.pack(fill=tk.X, padx=10)
        for title, desc, cb in [
            ("📝 文本处理","大小写/去重/排序/统计", self._text_tools),
            ("🔣 编码转换","Base64/URL/Hex", self._encode_tools),
            ("⏱ 时间工具","时间戳/日期差/实时", self._time_tools),
            ("🔑 哈希计算","MD5/SHA1/SHA256", self._hash_tools),
        ]: self._create_card(row, title, desc, cb)

    def _text_tools(self):
        win = tk.Toplevel(self.root); win.title("文本处理"); win.geometry("620x520"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="输入:", font=("微软雅黑", 10), bg=self.colors['light']).pack(pady=5)
        inp = scrolledtext.ScrolledText(win, height=10, font=("Consolas", 10)); inp.pack(fill=tk.X, padx=15, pady=5)
        out = scrolledtext.ScrolledText(win, height=10, font=("Consolas", 10)); out.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _op(op):
            t = inp.get("1.0", tk.END).rstrip("\n")
            out.delete("1.0", tk.END)
            if op=="upper": out.insert("1.0", t.upper())
            elif op=="lower": out.insert("1.0", t.lower())
            elif op=="blank": out.insert("1.0", "\n".join(l for l in t.split("\n") if l.strip()))
            elif op=="dedup":
                s=set(); r=[]
                for l in t.split("\n"):
                    if l not in s: s.add(l); r.append(l)
                out.insert("1.0", "\n".join(r))
            elif op=="sort": out.insert("1.0", "\n".join(sorted(t.split("\n"))))
            elif op=="stats":
                wc=len(t.split()); lc=len([l for l in t.split("\n") if l.strip()])
                out.insert("1.0", f"字符数: {len(t)}\n有效行: {lc}\n词数: {wc}")
        bf = tk.Frame(win, bg=self.colors['light']); bf.pack(pady=5)
        for o,l in [("upper","大写"),("lower","小写"),("blank","去空行"),("dedup","去重"),("sort","排序"),("stats","统计")]:
            tk.Button(bf, text=l, command=lambda o=o: _op(o), cursor="hand2", font=("微软雅黑", 9),
                     bg=self.colors['primary'], fg="white", bd=0).pack(side=tk.LEFT, padx=2)

    def _encode_tools(self):
        win = tk.Toplevel(self.root); win.title("编码转换"); win.geometry("620x480"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="输入:", font=("微软雅黑", 10), bg=self.colors['light']).pack(pady=5)
        inp = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 10)); inp.pack(fill=tk.X, padx=15, pady=5)
        out = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 10)); out.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _op(op):
            t = inp.get("1.0", tk.END).rstrip("\n")
            out.delete("1.0", tk.END)
            try:
                if op=="b64e": out.insert("1.0", base64.b64encode(t.encode()).decode())
                elif op=="b64d": out.insert("1.0", base64.b64decode(t).decode("utf-8", errors="replace"))
                elif op=="urle": out.insert("1.0", urllib.parse.quote(t))
                elif op=="urld": out.insert("1.0", urllib.parse.unquote(t))
                elif op=="hex": out.insert("1.0", t.encode().hex())
                elif op=="dehex":
                    try: out.insert("1.0", bytes.fromhex(t).decode("utf-8", errors="replace"))
                    except: out.insert("1.0", "无效十六进制")
            except Exception as e: out.insert("1.0", f"错误: {e}")
        bf = tk.Frame(win, bg=self.colors['light']); bf.pack(pady=5)
        for o,l in [("b64e","Base64编码"),("b64d","Base64解码"),("urle","URL编码"),("urld","URL解码"),("hex","UTF8→Hex"),("dehex","Hex→UTF8")]:
            tk.Button(bf, text=l, command=lambda o=o: _op(o), cursor="hand2", font=("微软雅黑", 9),
                     bg=self.colors['primary'], fg="white", bd=0).pack(side=tk.LEFT, padx=2)

    def _time_tools(self):
        win = tk.Toplevel(self.root); win.title("时间工具"); win.geometry("500x420"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="时间戳→日期", font=("微软雅黑", 11, "bold"), bg=self.colors['light']).pack(pady=10)
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack()
        ts_e = tk.Entry(f1, width=20, font=("微软雅黑", 10)); ts_e.pack(side=tk.LEFT, padx=5)
        ts_l = tk.Label(win, text="", font=("微软雅黑", 11), fg=self.colors['primary'], bg=self.colors['light']); ts_l.pack(pady=5)
        def _ts():
            try:
                ts=float(ts_e.get())
                if ts>1e12: ts/=1000
                ts_l.config(text=datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S"))
            except: ts_l.config(text="错误",fg="red")
        tk.Button(f1, text="转换", command=_ts, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 9)).pack(side=tk.LEFT)
        # 当前时间
        tk.Label(win, text="当前时间戳(实时)", font=("微软雅黑", 11, "bold"), bg=self.colors['light']).pack(pady=(20,5))
        nl = tk.Label(win, text="", font=("Consolas", 14, "bold"), fg=self.colors['primary'], bg=self.colors['light']); nl.pack()
        def _up(): nl.config(text=str(int(time.time()))); win.after(1000, _up)
        _up()
        # 日期差
        tk.Label(win, text="日期差", font=("微软雅黑", 11, "bold"), bg=self.colors['light']).pack(pady=(20,5))
        f2 = tk.Frame(win, bg=self.colors['light']); f2.pack()
        d1=tk.Entry(f2,width=12,font=("微软雅黑",10)); d1.pack(side=tk.LEFT,padx=3); d1.insert(0,"2026-01-01")
        d2=tk.Entry(f2,width=12,font=("微软雅黑",10)); d2.pack(side=tk.LEFT,padx=3); d2.insert(0,datetime.now().strftime("%Y-%m-%d"))
        dl = tk.Label(win, text="", font=("微软雅黑", 11), fg=self.colors['primary'], bg=self.colors['light']); dl.pack(pady=5)
        def _diff():
            try:
                dd=(datetime.strptime(d2.get(),"%Y-%m-%d")-datetime.strptime(d1.get(),"%Y-%m-%d")).days
                dl.config(text=f"相差 {abs(dd)} 天")
            except: dl.config(text="格式错误",fg="red")
        tk.Button(f2, text="计算", command=_diff, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 9)).pack(side=tk.LEFT,padx=5)

    def _hash_tools(self):
        win = tk.Toplevel(self.root); win.title("哈希计算"); win.geometry("560x380"); win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack(pady=15)
        fpv = tk.StringVar(); tk.Entry(f1, textvariable=fpv, width=40, font=("微软雅黑", 10)).pack(side=tk.LEFT, padx=5)
        tk.Button(f1, text="浏览", command=lambda: (f:=filedialog.askopenfilename(), fpv.set(f) if f else None),
                 cursor="hand2", font=("微软雅黑", 10)).pack(side=tk.LEFT)
        te = tk.Entry(win, width=50, font=("Consolas", 10)); te.pack(pady=5)
        tk.Label(win, text="↑ 或输入文本", font=("微软雅黑", 9), fg=self.colors['gray'], bg=self.colors['light']).pack()
        rl = tk.Label(win, text="", font=("Consolas", 10), justify=tk.LEFT, bg=self.colors['light']); rl.pack(pady=15)
        def _calc():
            data = b""
            if fpv.get() and Path(fpv.get()).is_file():
                data = Path(fpv.get()).read_bytes()
            elif te.get(): data = te.get().encode("utf-8")
            else: return messagebox.showwarning("提示","请选择文件或输入文本")
            try:
                out=""
                for a in ["md5","sha1","sha256","sha512"]:
                    out += f"{a.upper()}: {hashlib.new(a, data).hexdigest()}\n"
                rl.config(text=out)
            except Exception as e: rl.config(text=f"错误: {e}")
        tk.Button(win, text="🔑 计算哈希", command=_calc, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"), width=14).pack(pady=10)


    # ==================== 剪贴板 ====================
    def _show_clipboard_tools(self):
        self.clear_content()
        self._section_header("剪贴板历史", "捕获 · 回溯 · 一键粘贴")
        btn_row = tk.Frame(self.content_frame, bg=self.colors['light']); btn_row.pack(fill=tk.X, padx=20, pady=5)
        tk.Button(btn_row, text="📋 捕获剪贴板", command=self._capture_clip,
                 bg=self.colors['primary'], fg="white", font=("微软雅黑", 10), cursor="hand2", bd=0).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_row, text="🗑 清空", command=self._clear_clip,
                 bg=self.colors['danger'], fg="white", font=("微软雅黑", 10), cursor="hand2", bd=0).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_row, text="🔄 刷新", command=self._show_clipboard_tools,
                 font=("微软雅黑", 10), cursor="hand2").pack(side=tk.LEFT, padx=5)
        self.clipboard_listbox = tk.Listbox(self.content_frame, font=("Consolas", 10), height=16,
                                            bg="white", selectbackground=self.colors['primary'])
        self.clipboard_listbox.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        self.clipboard_listbox.bind("<Double-1>", self._paste_clip)
        self._refresh_clip_list()


    def _auto_clip_timer(self):
        """剪贴板自动捕获（每3秒）"""
        try:
            t = self.root.clipboard_get()
            if t and len(t) > 2 and t not in self.clipboard_history:
                self.clipboard_history.insert(0, t[:800])
                if len(self.clipboard_history) > 50:
                    self.clipboard_history.pop()
                self._refresh_clip_list()
        except:
            pass
        self.root.after(3000, self._auto_clip_timer)

    def _capture_clip(self):
        try:
            t = self.root.clipboard_get()
            if t and t not in self.clipboard_history:
                self.clipboard_history.insert(0, t[:800])
                if len(self.clipboard_history) > 50: self.clipboard_history.pop()
            self._refresh_clip_list()
        except Exception: messagebox.showinfo("提示", "剪贴板空/不可读")

    def _clear_clip(self):
        self.clipboard_history.clear(); self._refresh_clip_list()

    def _refresh_clip_list(self):
        if hasattr(self, "clipboard_listbox") and self.clipboard_listbox is not None:
            try:
                self.clipboard_listbox.delete(0, tk.END)
                for i, t in enumerate(self.clipboard_history[:50]):
                    d = t[:70].replace("\n", " ").replace("\t", " ")
                    self.clipboard_listbox.insert(tk.END, f"#{i+1}: {d}")
            except Exception: pass

    def _paste_clip(self, event):
        sel = self.clipboard_listbox.curselection()
        if sel and sel[0] < len(self.clipboard_history):
            self.root.clipboard_clear(); self.root.clipboard_append(self.clipboard_history[sel[0]])
            messagebox.showinfo("提示", "已复制到剪贴板")

    # ==================== 内容搜索 ====================
    def _show_search_tools(self):
        self.clear_content()
        self._section_header("文件内容搜索", "按关键词+类型递归搜索")
        sf = tk.Frame(self.content_frame, bg=self.colors['light']); sf.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(sf, text="目录:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=0, col=0, sticky="w")
        d_var = tk.StringVar(); tk.Entry(sf, textvariable=d_var, width=40, font=("微软雅黑", 10)).grid(row=0, col=1, padx=5)
        tk.Button(sf, text="浏览", command=lambda: (d:=filedialog.askdirectory(), d_var.set(d) if d else None),
                 cursor="hand2", font=("微软雅黑", 10)).grid(row=0, col=2)
        tk.Label(sf, text="关键词:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=1, col=0, sticky="w", pady=(10,0))
        kw_var = tk.StringVar(); tk.Entry(sf, textvariable=kw_var, width=40, font=("微软雅黑", 10)).grid(row=1, col=1, padx=5, pady=(10,0))
        tk.Label(sf, text="类型:", font=("微软雅黑", 10), bg=self.colors['light']).grid(row=2, col=0, sticky="w", pady=(10,0))
        ft_var = tk.StringVar(value="*.txt")
        ttk.Combobox(sf, textvariable=ft_var, values=["*.txt","*.py","*.md","*.csv","*.log","*.html","*.json","*.yaml","*.xml","*.*"],
                    state="readonly", width=16).grid(row=2, col=1, sticky="w", padx=5, pady=(10,0))
        # 按钮行
        btn_row = tk.Frame(self.content_frame, bg=self.colors['light'])
        btn_row.pack(fill=tk.X, padx=20, pady=(5, 0))
        self._search_cancel = False
        res = scrolledtext.ScrolledText(self.content_frame, height=16, font=("Consolas", 9), bg="white"); res.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        search_btn = tk.Button(btn_row, text="🔍 搜索", font=("微软雅黑", 10),
                 bg=self.colors['primary'], fg="white", cursor="hand2")
        def _search():
            d = d_var.get(); kw = kw_var.get().strip(); ft = ft_var.get()
            if not d or not kw: return messagebox.showwarning("提示", "请填写目录和关键词")
            res.delete("1.0", tk.END); self.set_status("搜索中…"); self._search_cancel = False
            search_btn.config(text="⏹ 停止", command=_cancel)
            def _cancel():
                self._search_cancel = True
                search_btn.config(text="🔍 搜索", command=_search)
            def _do_search():
                cnt, matches = 0, 0
                try:
                    for p in Path(d).rglob(ft):
                        if self._search_cancel: break
                        if p.is_file() and matches < 60:
                            cnt += 1
                            try:
                                content = p.read_text(encoding="utf-8", errors="ignore")
                                if kw.lower() in content.lower():
                                    matches += 1
                                    self.root.after(0, lambda p=p: res.insert(tk.END, f">> {p}\n"))
                                    for i, line in enumerate(content.split("\n"), 1):
                                        if kw.lower() in line.lower():
                                            self.root.after(0, lambda l=line,i=i: res.insert(tk.END, f"   L{i}: {l.strip()[:90]}\n"))
                                    self.root.after(0, lambda: res.insert(tk.END, "\n"))
                            except Exception: pass
                except Exception as e:
                    self.root.after(0, lambda e=e: res.insert(tk.END, f"错误: {e}\n"))
                status = "已取消" if self._search_cancel else "就绪"
                self.root.after(0, lambda: res.insert(tk.END, f"\n扫描: {cnt} 个  匹配: {matches} 个  [{status}]\n"))
                self.root.after(0, lambda: self.set_status(status))
                self.root.after(0, lambda: search_btn.config(text="🔍 搜索", command=_search))
                self._search_cancel = False
            threading.Thread(target=_do_search, daemon=True).start()
        search_btn.config(command=_search)
        search_btn.grid(row=2, column=2, pady=(10,0))

    # ==================== 日程管理 ====================
    def _show_calendar_tools(self):
        self.clear_content()
        self._section_header("日程管理", "待办事项 · 备忘录 (自动保存)")
        tf = DATA_DIR / "todo.json"; mf = DATA_DIR / "memo.txt"
        if tf.exists():
            try: self._todo_data = json.loads(tf.read_text(encoding="utf-8"))
            except Exception: self._todo_data = []
        else: self._todo_data = []
        # 待办 - 左侧
        left = tk.LabelFrame(self.content_frame, text="  📋 待办事项 ", font=("微软雅黑", 11, "bold"),
                             bg="white", padx=10, pady=5); left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(20,5), pady=5)
        inp = tk.Entry(left, font=("微软雅黑", 10)); inp.pack(fill=tk.X, pady=5)
        self.todo_listbox = tk.Listbox(left, font=("微软雅黑", 10), height=14, selectmode=tk.SINGLE, bg="white")
        self.todo_listbox.pack(fill=tk.BOTH, expand=True, pady=5)
        for item in self._todo_data:
            prefix = "✅ " if item.get("done") else "⬜ "
            self.todo_listbox.insert(tk.END, prefix + item.get("text", ""))
        bf = tk.Frame(left, bg="white"); bf.pack(fill=tk.X, pady=5)
        def _add():
            t = inp.get().strip()
            if not t: return
            self._todo_data.append({"text": t, "done": False, "time": datetime.now().isoformat()})
            self.todo_listbox.insert(tk.END, "⬜ " + t); inp.delete(0, tk.END); self._save_todo()
        def _tog():
            sel = self.todo_listbox.curselection()
            if not sel or sel[0] >= len(self._todo_data): return
            idx = sel[0]; self._todo_data[idx]["done"] = not self._todo_data[idx]["done"]
            done = self._todo_data[idx]["done"]
            pref = "✅ " if done else "⬜ "; self.todo_listbox.delete(idx)
            self.todo_listbox.insert(idx, pref + self._todo_data[idx]["text"]); self._save_todo()
        def _del():
            sel = self.todo_listbox.curselection()
            if not sel or sel[0] >= len(self._todo_data): return
            idx = sel[0]; self._todo_data.pop(idx); self.todo_listbox.delete(idx); self._save_todo()
        tk.Button(bf, text="添加", command=_add, cursor="hand2", bg=self.colors['primary'], fg="white",
                 bd=0, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=3)
        tk.Button(bf, text="切换", command=_tog, cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=3)
        tk.Button(bf, text="删除", command=_del, cursor="hand2", bg=self.colors['danger'], fg="white",
                 bd=0, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=3)
        inp.bind("<Return>", lambda e: _add())
        # 备忘录 - 右侧
        right = tk.LabelFrame(self.content_frame, text="  📝 备忘录 ", font=("微软雅黑", 11, "bold"),
                              bg="white", padx=10, pady=5); right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5,20), pady=5)
        self.memo_text = scrolledtext.ScrolledText(right, font=("微软雅黑", 10), height=14, bg="white")
        self.memo_text.pack(fill=tk.BOTH, expand=True, pady=5)
        if mf.exists():
            try: self.memo_text.insert("1.0", mf.read_text(encoding="utf-8"))
            except Exception: pass
        def _save_memo():
            try:
                mf.write_text(self.memo_text.get("1.0", tk.END), encoding="utf-8")
                messagebox.showinfo("💾", "备忘录已保存")
            except Exception as e: messagebox.showerror("错误", str(e))
        tk.Button(right, text="💾 保存", command=_save_memo, cursor="hand2",
                 bg=self.colors['success'], fg="white", bd=0, font=("微软雅黑", 10)).pack(pady=5)

    def _save_todo(self):
        try:
            (DATA_DIR / "todo.json").write_text(
                json.dumps(self._todo_data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception: pass

    def _split_file_dlg(self):
        fp = filedialog.askopenfilename(title="选择要分割的文件")
        if not fp: return
        win = tk.Toplevel(self.root); win.title("文件分割"); win.geometry("450x380")
        win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"文件: {Path(fp).name}", font=("微软雅黑",10), bg=self.colors['light']).pack(pady=8)
        tk.Label(win, text="分割方式:", font=("微软雅黑",10), bg=self.colors['light']).pack()
        mode = tk.StringVar(value="size")
        ttk.Combobox(win, textvariable=mode, values=["按大小","按行数"], state="readonly", width=12).pack()
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack(pady=10)
        tk.Label(f1, text="每个包大小:", font=("微软雅黑",10), bg=self.colors['light']).grid(row=0,col=0)
        sv = tk.Entry(f1, font=("微软雅黑",10), width=10); sv.grid(row=0,col=1,padx=5)
        tk.Label(f1, text=" KB (0=按行)", font=("微软雅黑",9), bg=self.colors['light']).grid(row=0,col=2)
        out_dir = filedialog.askdirectory(title="输出目录")
        log = scrolledtext.ScrolledText(win, height=10, font=("Consolas",9)); log.pack(fill=tk.BOTH,expand=True,padx=15,pady=5)
        def _go():
            size_kb = int(sv.get() or "0"); out = out_dir or str(Path(fp).parent)
            log.delete("1.0", tk.END)
            try:
                stem = Path(fp).stem; ext = Path(fp).suffix
                if mode.get()=="按行":
                    with open(fp, 'r', encoding='utf-8', errors='ignore') as f:
                        lines = f.readlines()
                    per = max(1, int(sv.get() or "500"))
                    for i in range(0, len(lines), per):
                        chunk = lines[i:i+per]
                        op = Path(out)/f"{stem}_p{i//per+1}{ext}"
                        op.write_text("".join(chunk), encoding='utf-8')
                        log.insert(tk.END, f"[OK] {op.name} ({len(chunk)}行)\n")
                else:
                    size_b = size_kb*1024
                    if size_b<=0: log.insert(tk.END,"请输入大于0的KB值\n"); return
                    with open(fp, 'rb') as f:
                        part = 1
                        while True:
                            data = f.read(size_b)
                            if not data: break
                            op = Path(out)/f"{stem}_p{part:03d}{ext}"
                            op.write_bytes(data)
                            log.insert(tk.END, f"[OK] {op.name} ({len(data)}B)\n")
                            part += 1
                log.insert(tk.END, f"\n✅ 分割完成，共 {part-1} 个文件\n")
            except Exception as e: log.insert(tk.END, f"错误: {e}\n")
        tk.Button(win, text="▶ 开始分割", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=14).pack(pady=8)

    def _archive_dlg(self):
        src = filedialog.askdirectory(title="选择文件夹打包")
        if not src: return
        save = filedialog.asksaveasfilename(title="保存压缩包",
            defaultextension=".zip", filetypes=[("ZIP","*.zip"),("7Z","*.7z")])
        if not save: return
        fmt = "zip" if save.endswith(".zip") else "gztar"
        def _go():
            import zipfile
            base = Path(src).name; log_text = []
            try:
                with zipfile.ZipFile(save, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for f in Path(src).rglob("*"):
                        if f.is_file():
                            arcname = Path(src).name + "/" + f.relative_to(Path(src))
                            zf.write(f, arcname)
                            log_text.append(f"  + {arcname}")
                self.root.after(0, lambda: self._show_success_dialog("完成", f"打包完成\n{save}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
        self._run_thread(_go, done_msg="打包完成")

    def _shortcut_dlg(self):
        win = tk.Toplevel(self.root); win.title("快捷方式管理"); win.geometry("520x420")
        win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="选择程序/文件:", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=10)
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack(fill=tk.X, padx=20)
        fv = tk.StringVar(); tk.Entry(f1, textvariable=fv, width=40, font=("微软雅黑",10)).pack(side=tk.LEFT, padx=5)
        tk.Button(f1, text="浏览", command=lambda: (d:=filedialog.askopenfilename(), fv.set(d) if d else None),
                 cursor="hand2", font=("微软雅黑",9)).pack(side=tk.LEFT)
        tk.Label(win, text="快捷方式名称:", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=8)
        nv = tk.Entry(win, width=30, font=("微软雅黑",11)); nv.pack()
        tk.Label(win, text="放置位置:", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=8)
        lv = tk.BooleanVar(value=True)
        tk.Checkbutton(win, text="桌面", variable=lv, font=("微软雅黑",10), bg=self.colors['light']).pack()
        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas",9)); log.pack(fill=tk.BOTH,expand=True,padx=15,pady=5)
        def _go():
            fp = fv.get().strip(); name = nv.get().strip() or Path(fp).stem
            if not fp or not Path(fp).exists(): log.insert(tk.END,"请选择有效文件\n"); return
            dest_dir = Path.home()/"Desktop" if lv.get() else Path.home()/"AppData"/"Roaming"/"Microsoft"/"Windows"/"Start Menu"/"Programs"
            try:
                import winshell
                winshell.CreateShortcut(
                    Path=dest_dir/f"{name}.lnk",
                    TargetPath=fp,
                    StartIn=str(Path(fp).parent)
                )
                log.insert(tk.END, f"[OK] 已创建: {dest_dir}/{name}.lnk\n")
            except Exception:
                # Windows API fallback
                try:
                    import pythoncom, win32com.client
                    pythoncom.CoInitialize()
                    shell = win32com.client.Dispatch("WScript.Shell")
                    sc = shell.CreateShortcut(str(dest_dir/f"{name}.lnk"))
                    sc.TargetPath = fp
                    sc.Save()
                    log.insert(tk.END, f"[OK] 已创建: {dest_dir}/{name}.lnk\n")
                except Exception as e:
                    log.insert(tk.END, f"[FAIL] {e}\n请以管理员权限运行\n")
        tk.Button(win, text="▶ 创建快捷方式", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=18).pack(pady=10)

    def _checksum_dlg(self):
        files = filedialog.askopenfilenames(title="选择文件(可多选)")
        if not files: return
        alg = tk.StringVar(value="MD5")
        win = tk.Toplevel(self.root); win.title("文件校验"); win.geometry("620x420")
        win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"已选 {len(files)} 个文件", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=5)
        ttk.Combobox(win, textvariable=alg, values=["MD5","SHA1","SHA256","SHA512"], state="readonly", width=10).pack()
        lb = scrolledtext.ScrolledText(win, font=("Consolas", 10), height=14); lb.pack(fill=tk.BOTH,expand=True,padx=15,pady=5)
        def _go():
            hfun = lambda d: hashlib.new(alg.get().lower(), d).hexdigest()
            lb.delete("1.0", tk.END)
            for fp in files:
                try:
                    data = Path(fp).read_bytes()
                    cs = hfun(data)
                    lb.insert(tk.END, f"{alg.get()}: {cs}  {Path(fp).name}\n")
                except Exception as e:
                    lb.insert(tk.END, f"[ERR] {Path(fp).name}: {e}\n")
        tk.Button(win, text="▶ 计算校验和", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=16).pack(pady=8)

    def _clean_dlg(self):
        win = tk.Toplevel(self.root); win.title("系统清理"); win.geometry("560x480")
        win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        tk.Label(win, text="清理临时文件和系统缓存", font=("微软雅黑",13,"bold"), bg=self.colors['light']).pack(pady=10)
        paths = [
            (Path.home()/"AppData"/"Local"/"Temp", "用户临时目录"),
            (Path("C:/Windows/Temp"), "系统临时目录"),
            (Path.home()/"AppData"/"Local"/"pip"/"Cache", "pip缓存"),
        ]
        vars_dict = {}
        for p, desc in paths:
            v = tk.BooleanVar(value=True)
            tk.Checkbutton(win, text=f"{desc}\n  {p}", variable=v, font=("微软雅黑",9),
                         bg=self.colors['light'], justify=tk.LEFT).pack(anchor="w", padx=30)
            vars_dict[p] = v
        log = scrolledtext.ScrolledText(win, height=10, font=("Consolas",9)); log.pack(fill=tk.BOTH,expand=True,padx=15,pady=5)
        def _do_clean():
            total = 0; count = 0
            for p, v in vars_dict.items():
                if not v.get(): continue
                self.root.after(0, lambda p=p: log.insert(tk.END, f"清理 {p} ...\n"))
                try:
                    for f in p.rglob("*"):
                        try:
                            if f.is_file():
                                sz = f.stat().st_size
                                f.unlink(); total += sz; count += 1
                        except Exception: pass
                    self.root.after(0, lambda c=count, t=total: log.insert(tk.END, f"  ✓ {c}个文件, 释放 {t//1024//1024}MB\n"))
                except Exception as e:
                    self.root.after(0, lambda: log.insert(tk.END, f"  ✗ {e}\n"))
            self.root.after(0, lambda t=total, c=count: log.insert(tk.END, f"\n✅ 共释放 {t//1024//1024} MB ({c}个文件)\n"))
        tk.Button(win, text="🧹 开始清理", command=lambda: threading.Thread(target=_do_clean, daemon=True).start(), cursor="hand2",
                 bg=self.colors['danger'], fg="white", font=("微软雅黑",11), width=16).pack(pady=8)

    def _diff_files_dlg(self):
        files = filedialog.askopenfilenames(title="选择两个文本文件对比", filetypes=[("文本","*.*")])
        if len(files) < 2: return messagebox.showinfo("提示","请选2个文件")
        f1, f2 = files[0], files[1]
        try:
            c1 = Path(f1).read_text(encoding="utf-8", errors="ignore")
            c2 = Path(f2).read_text(encoding="utf-8", errors="ignore")
        except Exception as e: return messagebox.showerror("错误", str(e))
        win = tk.Toplevel(self.root); win.title("文件对比"); win.geometry("900x600")
        win.transient(self.root); win.grab_set()
        win.configure(bg=self.colors['light'])
        # 左右分栏
        left = scrolledtext.ScrolledText(win, font=("Consolas", 10), bg="#FAFAFA")
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2, pady=5)
        right = scrolledtext.ScrolledText(win, font=("Consolas", 10), bg="#FAFAFA")
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2, pady=5)
        left.insert("1.0", c1); right.insert("1.0", c2)
        # 高亮差异
        lc = c1.split("\n"); rc = c2.split("\n")
        for i, (l, r) in enumerate(zip(lc, rc)):
            if l != r:
                left.tag_add(f"diff_{i}", f"{i+1}.0", f"{i+1}.end")
                right.tag_add(f"diff_{i}", f"{i+1}.0", f"{i+1}.end")
                left.tag_config(f"diff_{i}", background="#FFCCCC")
                right.tag_config(f"diff_{i}", background="#FFCCCC")
        # 行号列
        for i, line in enumerate(lc, 1):
            left.insert(f"{i}.0", f"{i:4d}| ")
        for i, line in enumerate(rc, 1):
            right.insert(f"{i}.0", f"{i:4d}| ")

    # ==================== Excel高级 ====================
    def _excel_chart_dlg(self):
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel","*.xlsx")])
        if not fp: return
        win = tk.Toplevel(self.root); win.title("生成图表"); win.geometry("500x400")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text="选择数据范围(行号):", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=10)
        f1 = tk.Frame(win, bg=self.colors['light']); f1.pack()
        tk.Label(f1, text="从第", font=("微软雅黑",10), bg=self.colors['light']).grid(row=0,col=0)
        sr = tk.Entry(f1, width=5, font=("微软雅黑",10)); sr.insert(0,"1"); sr.grid(row=0,col=1,padx=3)
        tk.Label(f1, text="行 到第", font=("微软雅黑",10), bg=self.colors['light']).grid(row=0,col=2)
        er = tk.Entry(f1, width=5, font=("微软雅黑",10)); er.insert(0,"20"); er.grid(row=0,col=3,padx=3)
        tk.Label(f1, text="行", font=("微软雅黑",10), bg=self.colors['light']).grid(row=0,col=4)
        tk.Label(win, text="图表类型:", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=8)
        ct = tk.StringVar(value="bar")
        ttk.Combobox(win, textvariable=ct, values=["bar","column","line","pie"], state="readonly", width=12).pack()
        save = filedialog.asksaveasfilename(title="保存图表", defaultextension=".xlsx", filetypes=[("Excel","*.xlsx")])
        def _go():
            try:
                s,e = int(sr.get()), int(er.get())
                wb = openpyxl.load_workbook(fp); ws = wb.active
                data = [[ws.cell(row=r,column=c).value for c in range(1, ws.max_column+1)] for r in range(s, e+1)]
                if not data: return messagebox.showwarning("提示","无数据")
                # 创建新工作簿放图表
                wb2 = openpyxl.Workbook(); ws2 = wb2.active; ws2.title = "数据源"
                for row_data in data: ws2.append(row_data)
                chart_type = {"bar":"bar","column":"column","line":"line","pie":"pie"}.get(ct.get(), "bar")
                from openpyxl.chart import BarChart, LineChart, PieChart, Reference
                if chart_type == "bar":
                    chart = BarChart(); chart.title = "图表"
                elif chart_type == "column":
                    from openpyxl.chart import BarChart
                    chart = BarChart(); chart.type = "col"
                elif chart_type == "line":
                    chart = LineChart()
                else:
                    chart = PieChart()
                cats = Reference(ws2, min_col=1, min_row=2, max_row=len(data))
                vals = Reference(ws2, min_col=2, min_row=1, max_row=len(data))
                chart.add_data(vals, titles_from_data=True)
                chart.set_categories(cats)
                ws2.add_chart(chart, "E2")
                wb2.save(save if save else fp.replace(".xlsx","_图表.xlsx")); wb2.close(); wb.close()
                self.root.after(0, lambda: self._show_success_dialog("完成", "图表已生成"))
            except Exception as ex:
                self.root.after(0, lambda: messagebox.showerror("错误", str(ex)))
        tk.Button(win, text="▶ 生成图表", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=14).pack(pady=10)

    def _vlookup_dlg(self):
        files = filedialog.askopenfilenames(title="选择两个Excel(第1个主表,第2个要匹配的表)", filetypes=[("Excel","*.xlsx")])
        if len(files)<2: return messagebox.showinfo("提示","请选2个Excel文件")
        save = filedialog.asksaveasfilename(title="保存结果", defaultextension=".xlsx", filetypes=[("Excel","*.xlsx")])
        if not save: return
        def _go():
            try:
                wb1 = openpyxl.load_workbook(files[0], read_only=True); ws1 = wb1.active
                wb2 = openpyxl.load_workbook(files[1], read_only=True); ws2 = wb2.active
                rows1 = list(ws1.iter_rows(values_only=True))
                rows2 = list(ws2.iter_rows(values_only=True))
                if not rows1 or not rows2: return messagebox.showwarning("提示","文件无数据")
                wb_out = openpyxl.Workbook(); ws_out = wb_out.active
                # 输出表头: 主表表头 + "匹配列"标签
                hdr = list(rows1[0]) + [f"{rows2[0][0]}_匹配"]
                ws_out.append(hdr)
                key_idx = 0  # 默认第0列
                # 建第2个表的索引
                key_map = {str(row[0]): row for row in rows2[1:]}
                for r in rows1[1:]:
                    key = str(r[key_idx]) if key_idx < len(r) else ""
                    match_row = key_map.get(key)
                    match_val = list(match_row) if match_row else [""]
                    ws_out.append(list(r) + match_val)
                wb_out.save(save); wb_out.close(); wb1.close(); wb2.close()
                self.root.after(0, lambda: self._show_success_dialog("完成", f"VLOOKUP完成\n{save}"))
            except Exception as ex:
                self.root.after(0, lambda: messagebox.showerror("错误", str(ex)))
        self._run_thread(_go, done_msg="VLOOKUP完成")

    def _conditional_format_dlg(self):
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel","*.xlsx")])
        if not fp: return
        win = tk.Toplevel(self.root); win.title("条件格式"); win.geometry("480x380")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text="选择列:", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=8)
        try:
            wb = openpyxl.load_workbook(fp, read_only=True); ws = wb.active
            rows = list(ws.iter_rows(values_only=True)); wb.close()
            if not rows: return messagebox.showwarning("提示","无数据")
            lb = tk.Listbox(win, height=8, font=("微软雅黑",10))
            for i,h in enumerate(rows[0]): lb.insert(tk.END, f"[{i}] {h}")
            lb.pack(fill=tk.X, padx=20)
        except Exception as e: return messagebox.showerror("错误", str(e))
        tk.Label(win, text="条件:", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=8)
        cond_v = tk.StringVar(value=">0")
        tk.Entry(win, textvariable=cond_v, width=20, font=("微软雅黑",10)).pack()
        def _go():
            sel = lb.curselection()
            if not sel: return messagebox.showwarning("提示","请选择列")
            ci = sel[0]; cond = cond_v.get()
            try:
                wb = openpyxl.load_workbook(fp); ws = wb.active
                for i, row in enumerate(ws.iter_rows(min_row=2), 2):
                    val = row[ci].value
                    if val is not None:
                        try:
                            if self._safe_cond_check(val, cond): 
                                for cell in row: cell.fill = PatternFill("solid", fgColor="FFFF00")
                        except: pass
                wb.save(fp); wb.close()
                self.root.after(0, lambda: self._show_success_dialog("完成","条件格式已应用"))
            except Exception as ex:
                self.root.after(0, lambda: messagebox.showerror("错误", str(ex)))
        tk.Button(win, text="▶ 应用格式", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=14).pack(pady=10)

    # ==================== 图片增强 ====================
    def _compress_img_dlg(self):
        files = filedialog.askopenfilenames(title="选择图片", filetypes=[("图片","*.png *.jpg *.jpeg *.bmp")])
        if not files: return
        win = tk.Toplevel(self.root); win.title("图片压缩"); win.geometry("450x320")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"已选 {len(files)} 张", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=8)
        tk.Label(win, text="压缩质量(1-100):", font=("微软雅黑",10), bg=self.colors['light']).pack()
        qv = tk.IntVar(value=75)
        tk.Scale(win, from_=10, to=100, orient="horizontal", variable=qv,
                font=("微软雅黑",10)).pack()
        tk.Label(win, text="输出格式:", font=("微软雅黑",10), bg=self.colors['light']).pack(pady=5)
        fmt = tk.StringVar(value="JPEG")
        ttk.Combobox(win, textvariable=fmt, values=["JPEG","PNG","WEBP"], state="readonly", width=10).pack()
        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas",9)); log.pack(fill=tk.BOTH,expand=True,padx=15,pady=5)
        def _go():
            q = qv.get(); f = fmt.get(); ok=fail=0; log.delete("1.0", tk.END)
            for fp in files:
                try:
                    img = Image.open(fp)
                    if img.mode == "RGBA" and f == "JPEG": img = img.convert("RGB")
                    out = fp.parent / f"{fp.stem}_compressed.{f.lower()}"
                    img.save(out, quality=q, optimize=True)
                    orig = fp.stat().st_size // 1024; new = out.stat().st_size // 1024
                    log.insert(tk.END, f"[OK] {fp.name}: {orig}KB → {new}KB (压缩{100*(orig-new)//orig if orig else 0}%)\n")
                    ok+=1
                except Exception as e: fail+=1; log.insert(tk.END, f"[FAIL] {fp.name}: {e}\n")
            log.insert(tk.END, f"\n✅ {ok}成功  ❌ {fail}失败\n")
        tk.Button(win, text="▶ 压缩", command=_go, cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=12).pack(pady=8)

    def _ocr_img_dlg(self):
        files = filedialog.askopenfilenames(title="选择图片", filetypes=[("图片","*.png *.jpg *.jpeg *.bmp")])
        if not files: return
        win = tk.Toplevel(self.root); win.title("图片OCR文字识别"); win.geometry("640x520")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"已选 {len(files)} 张图片", font=("微软雅黑",11), bg=self.colors['light']).pack(pady=5)
        result_text = scrolledtext.ScrolledText(win, font=("微软雅黑", 11), height=18); result_text.pack(fill=tk.BOTH,expand=True,padx=15,pady=5)
        def _do_ocr():
            try:
                import pytesseract
                for fp in files:
                    img = Image.open(fp)
                    text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    self.root.after(0, lambda t=text, f=fp: result_text.insert(tk.END, f"\n=== {Path(f).name} ===\n{t}\n"))
            except ImportError:
                self.root.after(0, lambda: result_text.insert(tk.END, "⚠️ OCR需要额外安装\n\n1. 下载 Tesseract-OCR:\n   https://github.com/UB-Mannheim/tesseract/wiki\n2. 安装后运行: pip install pytesseract\n3. 重启本软件即可使用OCR"))
            except Exception as e:
                self.root.after(0, lambda: result_text.insert(tk.END, f"[ERR] {e}\n"))
        def _save():
            save = filedialog.asksaveasfilename(title="保存文本", defaultextension=".txt", filetypes=[("Text","*.txt")])
            if save: Path(save).write_text(result_text.get("1.0", tk.END), encoding="utf-8")
        tk.Button(win, text="▶ 开始识别", command=lambda: threading.Thread(target=_do_ocr, daemon=True).start(), cursor="hand2",
                 bg=self.colors['primary'], fg="white", font=("微软雅黑",11), width=14).pack(pady=5)
        tk.Button(win, text="💾 保存文本", command=_save, cursor="hand2",
                 bg=self.colors['success'], fg="white", font=("微软雅黑",11), width=14).pack(pady=5)

    def _nine_grid_dlg(self):
        files = filedialog.askopenfilenames(title="选择1张图片(将切成9宫格)",
            filetypes=[("图片","*.png *.jpg *.jpeg")])
        if not files: return
        fp = files[0]
        save_dir = filedialog.askdirectory(title="保存到") or str(Path(fp).parent)
        def _go():
            try:
                img = Image.open(fp)
                w, h = img.size; sw, sh = w//3, h//3
                for i in range(3):
                    for j in range(3):
                        tile = img.crop((j*sw, i*sh, (j+1)*sw, (i+1)*sh))
                        out = Path(save_dir)/f"{Path(fp).stem}_grid_{i*3+j+1}{Path(fp).suffix}"
                        tile.save(out)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"已切成9张保存到\n{save_dir}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"切图失败: {e}"))
        self._run_thread(_go, done_msg="九宫格完成")

    # ==================== Word增强 ====================
    def _merge_word_dlg(self):
        files = filedialog.askopenfilenames(title="选择Word文件", filetypes=[("Word","*.docx")])
        if len(files)<2: return messagebox.showinfo("提示","至少选2个")
        save = filedialog.asksaveasfilename(title="保存合并", defaultextension=".docx", filetypes=[("Word","*.docx")])
        if not save: return
        def _go():
            from docx import Document
            doc_out = Document()
            for fp in files:
                doc = Document(fp)
                for para in doc.paragraphs:
                    doc_out.add_paragraph(para.text)
                for table in doc.tables:
                    t = doc_out.add_table(rows=1, cols=len(table.columns))
                    for ri, row in enumerate(table.rows):
                        for ci, cell in enumerate(row.cells):
                            t.rows[ri+1].cells[ci].text = cell.text
                doc_out.add_page_break()
            doc_out.save(save)
            self.root.after(0, lambda: self._show_success_dialog("完成", f"合并 {len(files)} 个文件\n{save}"))
        self._run_thread(_go, done_msg="合并完成")

    def _diff_word_dlg(self):
        files = filedialog.askopenfilenames(title="选择两个Word对比", filetypes=[("Word","*.docx")])
        if len(files)<2: return messagebox.showinfo("提示","请选2个Word")
        if not DOCX_AVAILABLE: return
        try:
            doc1 = Document(files[0]); doc2 = Document(files[1])
            t1 = "\n".join(p.text for p in doc1.paragraphs)
            t2 = "\n".join(p.text for p in doc2.paragraphs)
        except Exception as e: return messagebox.showerror("错误", str(e))
        win = tk.Toplevel(self.root); win.title("Word对比"); win.geometry("900x580")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        left = scrolledtext.ScrolledText(win, font=("Consolas", 10), bg="#FAFAFA")
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2, pady=5)
        right = scrolledtext.ScrolledText(win, font=("Consolas", 10), bg="#FAFAFA")
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2, pady=5)
        left.insert("1.0", t1); right.insert("1.0", t2)
        for i,(l,r) in enumerate(zip(t1.split("\n"), t2.split("\n"))):
            if l!=r:
                left.tag_add(f"d{i}", f"{i+1}.0", f"{i+1}.end")
                right.tag_add(f"d{i}", f"{i+1}.0", f"{i+1}.end")
                left.tag_config(f"d{i}", background="#FFCCCC")
                right.tag_config(f"d{i}", background="#FFCCCC")

    # ============================================================
    # 许可证系统 - 1年授权 · 到期遮罩 · 激活码验证
    # ============================================================

    LICENSE_FILE = DATA_DIR / "license.dat"
    USED_CODES_FILE = DATA_DIR / "used_codes.json"
    TRIAL_DAYS = 7  # 试用天数（7天）
    _ACTIVATION_KEY = b"YTQJ2025_OFFICE_ASSISTANT_PRO"

    def _update_status_bar_license(self):
        """更新状态栏的许可证信息（委托到LicenseManager）"""
        self.license.update_status_bar_license()

    def _safe_cond_check(self, val, cond):
        """安全条件比较(委托到utils.safe_cond_check)"""
        return safe_cond_check(val, cond)

    # ==================== 格式互转 ====================
    def _show_convert_tools(self):
        self.clear_content()
        self._section_header("格式互转", "Word↔PDF · Excel↔PDF · 图片↔PDF · Excel↔CSV↔JSON · Excel→HTML")
        # 显示模块加载状态
        if not REPORTLAB_AVAILABLE or not PYPDFIUM_AVAILABLE:
            warn = ""
            if not REPORTLAB_AVAILABLE:
                warn += f"⚠️ reportlab未加载(Doc/Excel/图片→PDF不可用)\n{_REPORTLAB_ERROR}\n\n"
            if not PYPDFIUM_AVAILABLE:
                warn += f"⚠️ pypdfium2未加载(PDF→图片不可用)\n{_PYPDFIUM_ERROR}\n"
            tk.Label(self.content_frame, text=warn, fg="red", font=("微软雅黑", 10),
                    bg=self.colors['light'], justify=tk.LEFT).pack(pady=5)
        for rdata in [
            [("📝 Word→PDF", "Word文档转PDF", self._word_to_pdf_dlg),
             ("📊 Excel→PDF", "Excel表格转PDF", self._excel_to_pdf_dlg),
             ("🖼 图片→PDF", "多图合为一个PDF", self._images_to_pdf_dlg)],
            [("📄 PDF→图片", "PDF页面转图片", self._pdf_to_img_dlg),
             ("📊 Excel→CSV", "导出CSV(UTF-8)", self._excel_to_csv_dlg),
             ("📋 CSV→Excel", "CSV导入Excel", self._csv_to_excel_dlg)],
            [("📊 Excel→JSON", "表格转JSON", self._excel_to_json_dlg),
             ("📋 JSON→Excel", "JSON导入Excel", self._json_to_excel_dlg),
             ("📊 Excel→HTML", "表格转网页", self._excel_to_html_dlg)],
            [("📋 CSV↔JSON", "CSV/JSON互转", self._csv_json_dlg),
             ("📝 Word→HTML", "Word转网页", self._word_to_html_dlg),
             ("🖼 图片批量转", "多格式互转", self._convert_img_dlg)],
        ]:
            row = tk.Frame(self.content_frame, bg=self.colors['light']); row.pack(fill=tk.X, padx=10)
            for title, desc, cb in rdata:
                self._create_card(row, title, desc, cb)

    def _word_to_pdf_dlg(self):
        """Word → PDF (reportlab + docx)"""
        if not DOCX_AVAILABLE:
            return messagebox.showerror("错误", f"python-docx未安装\n{_DOCX_ERROR if not DOCX_AVAILABLE else ''}")
        if not REPORTLAB_AVAILABLE:
            return messagebox.showerror("错误", f"reportlab未安装(Word→PDF依赖此库)\n{_REPORTLAB_ERROR}")
        fp = filedialog.askopenfilename(title="选择Word文档", filetypes=[("Word", "*.docx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存PDF", defaultextension=".pdf",
                                             filetypes=[("PDF", "*.pdf")])
        if not save: return
        def _go():
            try:
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RLTable
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import mm
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                # 注册中文字体（.ttc需subfontIndex）
                font_registered = False
                for fp_font in ["C:/Windows/Fonts/simhei.ttf", "C:/Windows/Fonts/msyh.ttc",
                                "C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/simfang.ttf"]:
                    if Path(fp_font).exists():
                        try:
                            if fp_font.endswith('.ttc'):
                                pdfmetrics.registerFont(TTFont('Chinese', fp_font, subfontIndex=0))
                            else:
                                pdfmetrics.registerFont(TTFont('Chinese', fp_font))
                            font_registered = True
                            break
                        except: continue
                doc = Document(fp)
                pdf_doc = SimpleDocTemplate(save, pagesize=A4)
                styles = getSampleStyleSheet()
                if font_registered:
                    for s in styles.byName.values():
                        s.fontName = 'Chinese'
                # 转义XML特殊字符
                def _esc(t):
                    return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                story = []
                for p in doc.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        story.append(Paragraph(_esc(txt), styles['Normal']))
                        story.append(Spacer(1, 4*mm))
                for table in doc.tables:
                    tdata = []
                    for row in table.rows:
                        tdata.append([_esc(cell.text) for cell in row.cells])
                    if tdata and tdata[0]:
                        ncols = len(tdata[0])
                        col_w = min(45*mm, (A4[0]-40*mm)/max(ncols,1))
                        story.append(RLTable(tdata, colWidths=[col_w]*ncols))
                        story.append(Spacer(1, 6*mm))
                if not story:
                    story.append(Paragraph("(空文档)", styles['Normal']))
                pdf_doc.build(story)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"Word→PDF 转换成功\n{save}"))
            except Exception as e:
                import traceback
                err = traceback.format_exc()
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{err[-600:]}"))
        self._run_thread(_go, done_msg="Word→PDF 完成")

    def _excel_to_pdf_dlg(self):
        """Excel → PDF (reportlab + openpyxl)"""
        if not OPENPYL_AVAILABLE:
            return messagebox.showerror("错误", f"openpyxl未安装\n{_OPENPYL_ERROR if not OPENPYL_AVAILABLE else ''}")
        if not REPORTLAB_AVAILABLE:
            return messagebox.showerror("错误", f"reportlab未安装(Excel→PDF依赖此库)\n{_REPORTLAB_ERROR}")
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel", "*.xlsx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存PDF", defaultextension=".pdf",
                                             filetypes=[("PDF", "*.pdf")])
        if not save: return
        def _go():
            try:
                from reportlab.platypus import SimpleDocTemplate, Table as RLTable, Spacer
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.lib.units import mm
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                # 注册中文字体（.ttc需subfontIndex）
                for fp_font in ["C:/Windows/Fonts/simhei.ttf", "C:/Windows/Fonts/msyh.ttc",
                                "C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/simfang.ttf"]:
                    if Path(fp_font).exists():
                        try:
                            if fp_font.endswith('.ttc'):
                                pdfmetrics.registerFont(TTFont('Chinese', fp_font, subfontIndex=0))
                            else:
                                pdfmetrics.registerFont(TTFont('Chinese', fp_font))
                            break
                        except: continue
                wb = openpyxl.load_workbook(fp, read_only=True)
                ws = wb.active
                # 转义XML特殊字符
                def _esc(t):
                    return safe_str(t).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append([_esc(c) for c in row])
                wb.close()
                if not data or not data[0]:
                    self.root.after(0, lambda: messagebox.showwarning("提示", "Excel为空"))
                    return
                ncols = len(data[0])
                page_size = landscape(A4) if ncols > 6 else A4
                pdf_doc = SimpleDocTemplate(save, pagesize=page_size)
                col_w = min(50*mm, (page_size[0] - 40*mm) / max(ncols, 1))
                story = [RLTable(data, colWidths=[col_w]*ncols,
                                  style=[('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                                         ('BACKGROUND', (0,0), (-1,0), colors.lightblue),
                                         ('FONTNAME', (0,0), (-1,-1), 'Chinese' if 'Chinese' in pdfmetrics._fonts else 'Helvetica'),
                                         ('FONTSIZE', (0,0), (-1,-1), 8)])]
                pdf_doc.build(story)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"Excel→PDF 转换成功\n{save}"))
            except Exception as e:
                import traceback
                err = traceback.format_exc()
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{err[-600:]}"))
        self._run_thread(_go, done_msg="Excel→PDF 完成")

    def _images_to_pdf_dlg(self):
        """多张图片 → 一个PDF (PIL + reportlab)"""
        if not PIL_AVAILABLE:
            return messagebox.showerror("错误", f"Pillow未安装\n{_PIL_ERROR}")
        if not REPORTLAB_AVAILABLE:
            return messagebox.showerror("错误", f"reportlab未安装(图片→PDF依赖此库)\n{_REPORTLAB_ERROR}")
        files = filedialog.askopenfilenames(title="选择图片（可多选）",
            filetypes=[("图片", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.webp")])
        if not files: return
        save = filedialog.asksaveasfilename(title="保存PDF", defaultextension=".pdf",
                                             filetypes=[("PDF", "*.pdf")])
        if not save: return
        def _go():
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import mm
                from reportlab.pdfgen import canvas
                c = canvas.Canvas(save, pagesize=A4)
                pw, ph = A4
                tmp_files = []
                for i, fp in enumerate(files):
                    img = Image.open(fp)
                    iw, ih = img.size
                    # 缩放适配A4
                    ratio = min((pw-20*mm)/max(iw,1), (ph-20*mm)/max(ih,1))
                    nw, nh = iw*ratio, ih*ratio
                    x, y = (pw-nw)/2, (ph-nh)/2
                    # 转为RGB JPEG给reportlab使用
                    if img.mode in ('RGBA', 'P', 'LA'):
                        bg = Image.new('RGB', img.size, (255,255,255))
                        if img.mode == 'RGBA':
                            bg.paste(img, mask=img.split()[-1])
                        else:
                            bg.paste(img)
                        img = bg
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    tmp = Path.home() / ".office_assistant" / f"_tmp_img_{i}.jpg"
                    tmp.parent.mkdir(exist_ok=True)
                    img.save(str(tmp), "JPEG", quality=95)
                    tmp_files.append(tmp)
                    c.drawImage(str(tmp), x, y, nw, nh)
                    c.showPage()
                c.save()
                # 清理临时文件
                for t in tmp_files:
                    t.unlink(missing_ok=True)
                self.root.after(0, lambda: self._show_success_dialog("完成", f"图片→PDF 成功\n{len(files)}张图片\n{save}"))
            except Exception as e:
                import traceback
                err = traceback.format_exc()
                # 清理临时文件
                for i in range(len(files)):
                    t = Path.home() / ".office_assistant" / f"_tmp_img_{i}.jpg"
                    t.unlink(missing_ok=True)
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{err[-600:]}"))
        self._run_thread(_go, done_msg="图片→PDF 完成")

    def _pdf_to_img_dlg(self):
        """PDF → 图片 (pypdfium2)"""
        if not PYPDFIUM_AVAILABLE:
            return messagebox.showerror("错误", f"pypdfium2未安装(PDF→图片依赖此库)\n{_PYPDFIUM_ERROR}")
        fp = filedialog.askopenfilename(title="选择PDF", filetypes=[("PDF", "*.pdf")])
        if not fp: return
        out_dir = filedialog.askdirectory(title="图片输出目录")
        if not out_dir: return
        win = tk.Toplevel(self.root); win.title("PDF转图片"); win.geometry("420x300")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text=f"PDF: {Path(fp).name}", font=("微软雅黑", 10),
                bg=self.colors['light']).pack(pady=8)
        tk.Label(win, text="输出DPI:", font=("微软雅黑", 10), bg=self.colors['light']).pack()
        dpi_var = tk.IntVar(value=150)
        ttk.Spinbox(win, from_=72, to=600, textvariable=dpi_var, width=8).pack(pady=5)
        fmt_var = tk.StringVar(value="PNG")
        tk.Label(win, text="输出格式:", font=("微软雅黑", 10), bg=self.colors['light']).pack()
        ttk.Combobox(win, textvariable=fmt_var, values=["PNG","JPEG"], state="readonly", width=8).pack(pady=5)
        log = scrolledtext.ScrolledText(win, height=6, font=("Consolas", 9)); log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        def _go():
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(fp)
                fmt = fmt_var.get().lower(); dpi = dpi_var.get(); ok = 0
                log.insert(tk.END, f"共 {len(pdf)} 页, DPI={dpi}, 格式={fmt.upper()}\n")
                for i in range(len(pdf)):
                    page = pdf[i]
                    bitmap = page.render(scale=dpi/72)
                    pil_img = bitmap.to_pil()
                    out_path = Path(out_dir) / f"page_{i+1:03d}.{fmt}"
                    if fmt == "jpeg":
                        if pil_img.mode in ("RGBA","P","LA"):
                            bg = Image.new("RGB", pil_img.size, (255,255,255))
                            bg.paste(pil_img, mask=pil_img.split()[-1] if pil_img.mode=="RGBA" else None)
                            pil_img = bg
                        else:
                            pil_img = pil_img.convert("RGB") if pil_img.mode != "RGB" else pil_img
                    if fmt == "jpeg":
                        pil_img.save(str(out_path), quality=95)
                    else:
                        pil_img.save(str(out_path))
                    ok += 1
                    log.insert(tk.END, f"  ✅ 第{i+1}页 → {out_path.name}\n")
                    log.see(tk.END); win.update()
                log.insert(tk.END, f"\n✅ 完成：{ok} 页已保存到 {out_dir}")
                self.set_status(f"PDF转图完成: {ok}页")
            except Exception as e:
                import traceback
                err = traceback.format_exc()
                log.insert(tk.END, f"\n❌ 错误:\n{err[-500:]}")
                self.root.after(0, lambda: messagebox.showerror("错误", f"PDF转图片失败:\n{err[-600:]}"))
        tk.Button(win, text="🔄 开始转换", command=lambda: self._run_thread(_go, done_msg="PDF转图完成"),
                 cursor="hand2", bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"),
                 width=14).pack(pady=8)

    def _csv_to_excel_dlg(self):
        """CSV → Excel"""
        if not OPENPYL_AVAILABLE:
            return messagebox.showerror("错误", "openpyxl未安装")
        fp = filedialog.askopenfilename(title="选择CSV", filetypes=[("CSV", "*.csv"), ("TSV", "*.tsv"), ("文本", "*.txt")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存Excel", defaultextension=".xlsx",
                                             filetypes=[("Excel", "*.xlsx")])
        if not save: return
        win = tk.Toplevel(self.root); win.title("CSV转Excel"); win.geometry("400x220")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text="编码格式:", font=("微软雅黑", 10), bg=self.colors['light']).pack(pady=8)
        enc_var = tk.StringVar(value="utf-8-sig")
        ttk.Combobox(win, textvariable=enc_var, values=["utf-8-sig","utf-8","gbk","gb18030","latin1"],
                     state="readonly", width=14).pack(pady=5)
        tk.Label(win, text="分隔符:", font=("微软雅黑", 10), bg=self.colors['light']).pack(pady=5)
        sep_var = tk.StringVar(value="auto")
        ttk.Combobox(win, textvariable=sep_var, values=["auto",",","\\t",";","|"],
                     state="readonly", width=14).pack(pady=5)
        def _go():
            import csv
            try:
                enc = enc_var.get()
                sep = sep_var.get()
                if sep == "auto":
                    # sniff delimiter
                    with open(fp, 'r', encoding=enc, errors='replace') as f:
                        sample = f.read(4096)
                    try:
                        sep = csv.Sniffer().sniff(sample).delimiter
                    except: sep = ","
                elif sep == "\\t":
                    sep = "\t"
                wb = openpyxl.Workbook(); ws = wb.active; ws.title = "数据"
                with open(fp, 'r', encoding=enc, errors='replace') as f:
                    reader = csv.reader(f, delimiter=sep)
                    for row in reader:
                        ws.append(row)
                wb.save(save); wb.close()
                row_count = ws.max_row
                self.root.after(0, lambda: self._show_success_dialog("完成", f"CSV→Excel 成功\n{row_count} 行\n{save}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{e}"))
        tk.Button(win, text="🔄 开始转换", command=lambda: self._run_thread(_go, done_msg="CSV→Excel完成"),
                 cursor="hand2", bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"),
                 width=14).pack(pady=12)

    def _excel_to_json_dlg(self):
        """Excel → JSON"""
        if not OPENPYL_AVAILABLE:
            return messagebox.showerror("错误", "openpyxl未安装")
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel", "*.xlsx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存JSON", defaultextension=".json",
                                             filetypes=[("JSON", "*.json")])
        if not save: return
        def _go():
            try:
                wb = openpyxl.load_workbook(fp, read_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(values_only=True))
                wb.close()
                if not rows: return
                headers = [safe_str(h) for h in rows[0]]
                data = []
                for row in rows[1:]:
                    item = {}
                    for i, val in enumerate(row):
                        key = headers[i] if i < len(headers) else f"col_{i}"
                        item[key] = safe_str(val) if val is not None else None
                    data.append(item)
                Path(save).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                self.root.after(0, lambda: self._show_success_dialog("完成", f"Excel→JSON 成功\n{len(data)} 条记录\n{save}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{e}"))
        self._run_thread(_go, done_msg="Excel→JSON 完成")

    def _json_to_excel_dlg(self):
        """JSON → Excel"""
        if not OPENPYL_AVAILABLE:
            return messagebox.showerror("错误", "openpyxl未安装")
        fp = filedialog.askopenfilename(title="选择JSON", filetypes=[("JSON", "*.json")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存Excel", defaultextension=".xlsx",
                                             filetypes=[("Excel", "*.xlsx")])
        if not save: return
        def _go():
            try:
                data = json.loads(Path(fp).read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    data = [data]
                if not data:
                    self.root.after(0, lambda: messagebox.showwarning("提示", "JSON为空"))
                    return
                wb = openpyxl.Workbook(); ws = wb.active; ws.title = "数据"
                headers = list(data[0].keys()) if data else []
                ws.append(headers)
                for item in data:
                    ws.append([safe_str(item.get(h, "")) for h in headers])
                wb.save(save); wb.close()
                self.root.after(0, lambda: self._show_success_dialog("完成", f"JSON→Excel 成功\n{len(data)} 条记录\n{save}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{e}"))
        self._run_thread(_go, done_msg="JSON→Excel 完成")

    def _excel_to_html_dlg(self):
        """Excel → HTML"""
        if not OPENPYL_AVAILABLE:
            return messagebox.showerror("错误", "openpyxl未安装")
        fp = filedialog.askopenfilename(title="选择Excel", filetypes=[("Excel", "*.xlsx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存HTML", defaultextension=".html",
                                             filetypes=[("HTML", "*.html")])
        if not save: return
        def _go():
            try:
                wb = openpyxl.load_workbook(fp, read_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(values_only=True))
                wb.close()
                html = ['<!DOCTYPE html><html><head><meta charset="utf-8">']
                html.append(f'<title>{Path(fp).stem}</title>')
                html.append('<style>body{font-family:微软雅黑,sans-serif;padding:20px;background:#f8f9fa}')
                html.append('table{border-collapse:collapse;width:100%}')
                html.append('th{background:#5B67CA;color:white;padding:8px 12px;text-align:left}')
                html.append('td{border:1px solid #dee2e6;padding:6px 10px}')
                html.append('tr:nth-child(even){background:#f2f2f2}')
                html.append('tr:hover{background:#e8ecf7}</style></head><body>')
                html.append(f'<h2>{Path(fp).stem}</h2><table>')
                for i, row in enumerate(rows):
                    tag = 'th' if i == 0 else 'td'
                    html.append('<tr>' + ''.join(f'<{tag}>{safe_str(c)}</{tag}>' for c in row) + '</tr>')
                html.append('</table></body></html>')
                Path(save).write_text('\n'.join(html), encoding="utf-8")
                self.root.after(0, lambda: self._show_success_dialog("完成", f"Excel→HTML 成功\n{save}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{e}"))
        self._run_thread(_go, done_msg="Excel→HTML 完成")

    def _csv_json_dlg(self):
        """CSV ↔ JSON 互转"""
        win = tk.Toplevel(self.root); win.title("CSV↔JSON互转"); win.geometry("500x380")
        win.transient(self.root); win.grab_set(); win.configure(bg=self.colors['light'])
        tk.Label(win, text="CSV ↔ JSON 格式互转", font=("微软雅黑", 13, "bold"),
                bg=self.colors['light']).pack(pady=10)
        mode_var = tk.StringVar(value="csv2json")
        mf = tk.Frame(win, bg=self.colors['light']); mf.pack(pady=5)
        for val, txt in [("csv2json", "CSV → JSON"), ("json2csv", "JSON → CSV")]:
            tk.Radiobutton(mf, text=txt, variable=mode_var, value=val,
                          font=("微软雅黑", 11), bg=self.colors['light']).pack(side=tk.LEFT, padx=15)
        # Encoding
        enc_frame = tk.Frame(win, bg=self.colors['light']); enc_frame.pack(pady=5)
        tk.Label(enc_frame, text="编码:", font=("微软雅黑", 10), bg=self.colors['light']).pack(side=tk.LEFT, padx=5)
        enc_var = tk.StringVar(value="utf-8-sig")
        ttk.Combobox(enc_frame, textvariable=enc_var, values=["utf-8-sig","utf-8","gbk","gb18030"],
                     state="readonly", width=10).pack(side=tk.LEFT)
        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=8)
        def _go():
            import csv
            mode = mode_var.get()
            enc = enc_var.get()
            try:
                if mode == "csv2json":
                    fp = filedialog.askopenfilename(title="选择CSV", filetypes=[("CSV", "*.csv")])
                    if not fp: return
                    save = filedialog.asksaveasfilename(title="保存JSON", defaultextension=".json",
                                                         filetypes=[("JSON", "*.json")])
                    if not save: return
                    with open(fp, 'r', encoding=enc, errors='replace') as f:
                        reader = csv.reader(f)
                        rows = list(reader)
                    if not rows: return
                    headers = rows[0]
                    data = [dict(zip(headers, [v if v else None for v in row])) for row in rows[1:]]
                    Path(save).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                    log.insert(tk.END, f"✅ CSV→JSON: {len(data)} 条记录已保存\n")
                else:
                    fp = filedialog.askopenfilename(title="选择JSON", filetypes=[("JSON", "*.json")])
                    if not fp: return
                    save = filedialog.asksaveasfilename(title="保存CSV", defaultextension=".csv",
                                                         filetypes=[("CSV", "*.csv")])
                    if not save: return
                    data = json.loads(Path(fp).read_text(encoding="utf-8"))
                    if isinstance(data, dict): data = [data]
                    if not data: return
                    headers = list(data[0].keys())
                    with open(save, 'w', encoding=enc, newline='') as f:
                        writer = csv.writer(f)
                        writer.writerow(headers)
                        for item in data:
                            writer.writerow([safe_str(item.get(h, "")) for h in headers])
                    log.insert(tk.END, f"✅ JSON→CSV: {len(data)} 条记录已保存\n")
            except Exception as e:
                log.insert(tk.END, f"❌ 错误: {e}\n")
        tk.Button(win, text="🔄 选择文件并转换", command=lambda: self._run_thread(_go, done_msg="互转完成"),
                 cursor="hand2", bg=self.colors['primary'], fg="white", font=("微软雅黑", 11, "bold"),
                 width=18).pack(pady=8)

    def _word_to_html_dlg(self):
        """Word → HTML"""
        if not DOCX_AVAILABLE:
            return messagebox.showerror("错误", "python-docx未安装")
        fp = filedialog.askopenfilename(title="选择Word", filetypes=[("Word", "*.docx")])
        if not fp: return
        save = filedialog.asksaveasfilename(title="保存HTML", defaultextension=".html",
                                             filetypes=[("HTML", "*.html")])
        if not save: return
        def _go():
            try:
                doc = Document(fp)
                html = ['<!DOCTYPE html><html><head><meta charset="utf-8">']
                html.append(f'<title>{Path(fp).stem}</title>')
                html.append('<style>body{font-family:微软雅黑,sans-serif;max-width:800px;margin:20px auto;padding:0 20px}')
                html.append('h1{color:#1E293B}h2{color:#334155}p{line-height:1.8}')
                html.append('table{border-collapse:collapse;width:100%}')
                html.append('th,td{border:1px solid #dee2e6;padding:8px 12px}')
                html.append('th{background:#5B67CA;color:white}</style></head><body>')
                for p in doc.paragraphs:
                    txt = p.text.strip()
                    if not txt: continue
                    if p.style.name.startswith('Heading 1'):
                        html.append(f'<h1>{txt}</h1>')
                    elif p.style.name.startswith('Heading'):
                        html.append(f'<h2>{txt}</h2>')
                    else:
                        html.append(f'<p>{txt}</p>')
                for table in doc.tables:
                    html.append('<table>')
                    for i, row in enumerate(table.rows):
                        tag = 'th' if i == 0 else 'td'
                        html.append('<tr>' + ''.join(f'<{tag}>{cell.text}</{tag}>' for cell in row.cells) + '</tr>')
                    html.append('</table>')
                html.append('</body></html>')
                Path(save).write_text('\n'.join(html), encoding="utf-8")
                self.root.after(0, lambda: self._show_success_dialog("完成", f"Word→HTML 成功\n{save}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"转换失败:\n{e}"))
        self._run_thread(_go, done_msg="Word→HTML 完成")

    def _about_dialog(self):
        """委托到LicenseManager"""
        self.license.about_dialog()


if __name__ == "__main__":
    app = OfficeAssistant()
    app.root.mainloop()