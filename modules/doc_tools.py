# -*- coding: utf-8 -*-
"""万能办公助手 v6.2 — DocTools（Word→文本·批量转换·表格提取·合并·对比）"""
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from pathlib import Path
import os, sys, threading
from datetime import datetime

from utils import DOCX_AVAILABLE


class DocToolsMixin:
    """DocTools — 所有方法通过 self 访问 OfficeAssistant 的属性"""

    def _show_doc_tools(self):
        self.clear_content()
        self._section_header("Word文档工具", "转换 · 批处理 · 表格提取 · 合并 · 对比")
        self._show_tips(
            "点击下方卡片 → 选择Word文件(.docx) → 自动处理 → 保存结果",
            "仅支持 .docx 格式，旧版 .doc 请先用Word另存为 .docx"
        )
        if not DOCX_AVAILABLE:
            tk.Label(
                self.content_frame,
                text="⚠️ python-docx 未安装\n请运行: pip install python-docx",
                fg="red", font=("微软雅黑", 12),
                bg=self.colors['light']
            ).pack(pady=60)
            return
        # 第一行 3 张卡片
        row1 = tk.Frame(self.content_frame, bg=self.colors['light'])
        row1.pack(fill=tk.X, padx=10)
        self._create_card(row1, "📝 Word→文本", "提取段落生成为TXT", self._word_to_text_dlg)
        self._create_card(row1, "📚 批量Word转", "批量提取多个Word文本", self._batch_word_dlg)
        self._create_card(row1, "📊 表格提取", "Word表格→Excel导出", self._extract_tables_dlg)
        # 第二行 2 张卡片（居中留空一侧均匀分布）
        row2 = tk.Frame(self.content_frame, bg=self.colors['light'])
        row2.pack(fill=tk.X, padx=10, pady=5)
        self._create_card(row2, "🔗 合并Word", "多个docx合并为一个", self._merge_word_dlg)
        self._create_card(row2, "🔄 对比Word", "逐行对比两个文档", self._diff_word_dlg)

    # ── Word→文本 ──────────────────────────────────────────────────────

    def _word_to_text_dlg(self):
        """选择单个 .docx → 提取所有段落 → 保存为 .txt"""
        fp = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not fp:
            return
        save = filedialog.asksaveasfilename(
            title="保存文本文件",
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt")]
        )
        if not save:
            return

        def _go():
            try:
                import docx
                doc = docx.Document(fp)
                lines = [p.text for p in doc.paragraphs]
                doc._body  # 确保已加载
                text = '\n'.join(lines)
                with open(save, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.root.after(0, lambda: self._show_success_dialog(
                    "完成",
                    f"提取完成\n{Path(fp).name} → {Path(save).name}\n共 {len(lines)} 段落"
                ))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))

        self._run_thread(_go, done_msg="Word→文本完成")

    # ── 批量Word转 ─────────────────────────────────────────────────────

    def _batch_word_dlg(self):
        """批量选择多个 .docx → 分别提取文本 → 保存到选择的目录"""
        files = filedialog.askopenfilenames(
            title="选择多个Word文档",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not files:
            return
        out_dir = filedialog.askdirectory(title="选择输出目录")
        if not out_dir:
            return

        def _go():
            try:
                import docx
                total = len(files)
                for i, fp in enumerate(files):
                    doc = docx.Document(fp)
                    text = '\n'.join(p.text for p in doc.paragraphs)
                    stem = Path(fp).stem
                    out_path = Path(out_dir) / f"{stem}.txt"
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                self.root.after(0, lambda: self._show_success_dialog(
                    "完成",
                    f"批量提取完成\n共 {total} 个文件\n输出目录: {out_dir}"
                ))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))

        self._run_thread(_go, done_msg="批量转换完成")

    # ── 表格提取 ───────────────────────────────────────────────────────

    def _extract_tables_dlg(self):
        """读取 Word 中的所有表格 → 每个表格写入 Excel（一个表格一个工作表）"""
        fp = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not fp:
            return
        save = filedialog.asksaveasfilename(
            title="保存Excel文件",
            defaultextension=".xlsx",
            filetypes=[("Excel 工作簿", "*.xlsx")]
        )
        if not save:
            return

        def _go():
            try:
                import docx
                doc = docx.Document(fp)
                tables = doc.tables
                if not tables:
                    self.root.after(0, lambda: messagebox.showwarning(
                        "提示", "文档中未发现表格"
                    ))
                    return

                from openpyxl import Workbook
                from openpyxl.styles import Alignment, Font
                wb = Workbook()
                # 删除默认工作表
                default_ws = wb.active
                wb.remove(default_ws) if default_ws is not None else None

                header_font = Font(bold=True)
                header_align = Alignment(horizontal='center', vertical='center')

                for idx, table in enumerate(tables):
                    ws_name = f"表格{idx+1}"[:31]  # Excel 工作表名最长31字符
                    ws = wb.create_sheet(title=ws_name)

                    for row_idx, row in enumerate(table.rows):
                        cells = [cell.text for cell in row.cells]
                        ws.append(cells)
                        # 首行加粗居中
                        if row_idx == 0:
                            for col_idx in range(len(cells)):
                                cell = ws.cell(row=row_idx+1, column=col_idx+1)
                                cell.font = header_font
                                cell.alignment = header_align

                    # 自动调整列宽（取前20列）
                    for col_idx in range(1, min(len(table.columns), 20) + 1):
                        max_len = 0
                        for row in ws.iter_rows(min_col=col_idx, max_col=col_idx,
                                                values_only=False):
                            for cell in row:
                                val = str(cell.value) if cell.value is not None else ""
                                # 中文字符按2个宽度算
                                cjk_count = sum(1 for c in val if '\u4e00' <= c <= '\u9fff')
                                length = len(val) + cjk_count
                                if length > max_len:
                                    max_len = length
                        ws.column_dimensions[chr(64 + col_idx)].width = min(max_len + 2, 50)

                wb.save(save)
                wb.close()
                self.root.after(0, lambda: self._show_success_dialog(
                    "完成",
                    f"表格提取完成\n共提取 {len(tables)} 个表格\n{save}"
                ))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))

        self._run_thread(_go, done_msg="表格提取完成")

    # ── 合并Word ───────────────────────────────────────────────────────

    def _merge_word_dlg(self):
        """合并多个 .docx 文件（保留段落结构，每个文档之间加分页符）"""
        files = filedialog.askopenfilenames(
            title="选择多个Word文档（按列表顺序合并）",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not files:
            return
        save = filedialog.asksaveasfilename(
            title="保存合并结果",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not save:
            return

        def _go():
            try:
                import docx
                from docx.enum.text import WD_BREAK

                # 以第一个文档为基础
                merged = docx.Document(files[0])

                for fp in files[1:]:
                    doc = docx.Document(fp)
                    # 在每个文档之前加一个分页符（第一个文档不加）
                    merged.add_page_break()
                    # 复制段落
                    for para in doc.paragraphs:
                        if para.text.strip() or para.runs:
                            new_para = merged.add_paragraph()
                            new_para.style = para.style
                            new_para.alignment = para.alignment
                            # 复制格式
                            for run in para.runs:
                                new_run = new_para.add_run(run.text, style=run.style)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                new_run.font.size = run.font.size
                                new_run.font.name = run.font.name
                                new_run.font.color.rgb = run.font.color.rgb

                merged.save(save)
                self.root.after(0, lambda: self._show_success_dialog(
                    "完成",
                    f"合并成功\n共 {len(files)} 个文档 → {Path(save).name}"
                ))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))

        self._run_thread(_go, done_msg="合并完成")

    # ── 对比Word ───────────────────────────────────────────────────────

    def _diff_word_dlg(self):
        """逐行对比两个 .docx 文档的文本内容"""
        # 选择第一个文档
        fp1 = filedialog.askopenfilename(
            title="选择第一个Word文档（旧版）",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not fp1:
            return
        # 选择第二个文档
        fp2 = filedialog.askopenfilename(
            title="选择第二个Word文档（新版）",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not fp2:
            return

        win = tk.Toplevel(self.root)
        win.title("文档对比结果")
        win.geometry("900x600")
        win.transient(self.root)
        win.grab_set()

        # 顶部状态栏
        status_frame = tk.Frame(win, bg="#F8FAFC")
        status_frame.pack(fill=tk.X, padx=10, pady=6)
        status_var = tk.StringVar(value="正在对比...")
        tk.Label(
            status_frame, textvariable=status_var,
            font=("微软雅黑", 10), bg="#F8FAFC", fg="#475569"
        ).pack(side=tk.LEFT)

        # 对比结果文本区域
        text_area = scrolledtext.ScrolledText(
            win, font=("Consolas", 10), wrap=tk.WORD,
            state=tk.DISABLED
        )
        text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))

        # 底部按钮
        btn_frame = tk.Frame(win, bg=win.cget('bg'))
        btn_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        tk.Button(
            btn_frame, text="关闭",
            command=win.destroy,
            bg=self.colors.get('primary', '#4F46E5'),
            fg="white", font=("微软雅黑", 10, "bold"),
            cursor="hand2", width=10
        ).pack(side=tk.RIGHT)

        def _go():
            try:
                import docx
                doc1 = docx.Document(fp1)
                doc2 = docx.Document(fp2)

                lines1 = [p.text for p in doc1.paragraphs]
                lines2 = [p.text for p in doc2.paragraphs]

                max_len = max(len(lines1), len(lines2))
                diff_lines = []
                same_count = 0
                diff_count = 0
                added_count = 0
                removed_count = 0

                for i in range(max_len):
                    l1 = lines1[i] if i < len(lines1) else None
                    l2 = lines2[i] if i < len(lines2) else None

                    line_num = i + 1
                    if l1 == l2:
                        # 相同行（灰色显示）
                        diff_lines.append(f"  {line_num:4d} | {l1}")
                        same_count += 1
                    elif l1 is None:
                        # 新增行
                        diff_lines.append(f"+ {line_num:4d} | {l2}")
                        added_count += 1
                    elif l2 is None:
                        # 删除行
                        diff_lines.append(f"- {line_num:4d} | {l1}")
                        removed_count += 1
                    else:
                        # 修改行
                        diff_lines.append(f"- {line_num:4d} | {l1}")
                        diff_lines.append(f"+ {line_num:4d} | {l2}")
                        diff_count += 1

                # 构建摘要
                summary = (
                    f"📊 对比完成\n"
                    f"  相同行: {same_count}  |  修改行: {diff_count}\n"
                    f"  新增行: {added_count}  |  删除行: {removed_count}\n"
                    f"  总行数: 文档1={len(lines1)}  文档2={len(lines2)}\n"
                    f"  ── 相同行(灰色) | -删除行 | +新增行 ──\n\n"
                )

                # 更新UI（线程安全）
                def _update_ui():
                    status_var.set(
                        f"✅ 对比完成 — 相同{same_count} 修改{diff_count} "
                        f"新增{added_count} 删除{removed_count}"
                    )
                    text_area.config(state=tk.NORMAL)
                    text_area.delete("1.0", tk.END)
                    text_area.insert(tk.END, summary, "summary")

                    # 分行插入，用不同颜色标记
                    for line in diff_lines:
                        if line.startswith("+ "):
                            text_area.insert(tk.END, line + "\n", "added")
                        elif line.startswith("- "):
                            text_area.insert(tk.END, line + "\n", "removed")
                        else:
                            text_area.insert(tk.END, line + "\n", "same")

                    text_area.config(state=tk.DISABLED)

                    # 配置标签样式
                    text_area.tag_config("summary", foreground="#1E293B",
                                         font=("微软雅黑", 10, "bold"))
                    text_area.tag_config("added", foreground="#16A34A")   # 绿色
                    text_area.tag_config("removed", foreground="#DC2626")  # 红色
                    text_area.tag_config("same", foreground="#94A3B8")    # 灰色

                self.root.after(0, _update_ui)

            except Exception as e:
                self.root.after(0, lambda: (
                    status_var.set(f"❌ 对比失败"),
                    messagebox.showerror("错误", str(e))
                ))

        self._run_thread(_go, done_msg="对比完成")
