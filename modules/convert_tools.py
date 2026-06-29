# -*- coding: utf-8 -*-
"""万能办公助手 — ConvertTools（格式互转：Word/Excel/PDF/图片/CSV/JSON/HTML）"""
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from pathlib import Path
import os, sys, json, csv, io, threading
from datetime import datetime

# ── 安全导入 utils ──────────────────────────────────────────────────────────
try:
    from utils import (
        DOCX_AVAILABLE, OPENPYL_AVAILABLE, PIL_AVAILABLE,
        PYPDFIUM_AVAILABLE, REPORTLAB_AVAILABLE, WORD_COM_AVAILABLE,
        register_chinese_font, get_chinese_font_name, esc_xml, safe_str,
    )
except ImportError:
    DOCX_AVAILABLE = False
    OPENPYL_AVAILABLE = False
    PIL_AVAILABLE = False
    PYPDFIUM_AVAILABLE = False
    REPORTLAB_AVAILABLE = False
    WORD_COM_AVAILABLE = False

    def register_chinese_font():
        pass

    def get_chinese_font_name() -> str:
        return "Helvetica"

    def esc_xml(text: str) -> str:
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")
        return text

    def safe_str(obj, default=""):
        try:
            return str(obj)
        except Exception:
            return default

# 错误消息变量 — 可能 utils 尚未定义，本地兜底
for _var_name in ("_DOCX_ERROR", "_OPENPYL_ERROR", "_PIL_ERROR",
                  "_PYPDFIUM_ERROR", "_REPORTLAB_ERROR"):
    try:
        exec(f"from utils import {_var_name}")
    except ImportError:
        vars()[_var_name] = ""


# ── 惰性导入第三方库（避免模块加载时报错）───────────────────────────────
def _import_docx():
    try:
        from docx import Document
        return Document
    except ImportError:
        return None

def _import_openpyxl():
    try:
        import openpyxl
        return openpyxl
    except ImportError:
        return None

def _import_pil():
    try:
        from PIL import Image
        return Image
    except ImportError:
        return None

def _import_pypdfium2():
    try:
        import pypdfium2 as pdfium
        return pdfium
    except ImportError:
        return None

def _import_reportlab():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
        from reportlab.lib.colors import HexColor, black
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        return (A4, mm, canvas, HexColor, black, SimpleDocTemplate, Paragraph,
                Spacer, RLImage, Table, TableStyle, getSampleStyleSheet,
                ParagraphStyle, TA_CENTER, TA_LEFT)
    except ImportError:
        return None

def _import_win32com():
    try:
        import win32com.client
        return win32com.client
    except ImportError:
        return None


class ConvertToolsMixin:
    """ConvertTools — 格式互转工具集（Word/Excel/PDF/图片/CSV/JSON/HTML）。"""

    # ══════════════════════════════════════════════════════════════════════
    #  主界面 — 3行4列卡片
    # ══════════════════════════════════════════════════════════════════════

    def _show_convert_tools(self):
        """显示格式互转主界面：3行4列卡片网格。"""
        self.clear_content()
        self._section_header("🔄 格式互转", "Word · Excel · PDF · 图片 · CSV · JSON · HTML 一键互转")
        self._show_tips(
            "点击卡片进入对应转换工具 → 选择源文件 → 设置参数 → 执行转换",
            "Word/Excel 转换需要安装对应库（python-docx / openpyxl），PDF 相关需要 reportlab/pypdfium2"
        )

        # 第一行：PDF 相关
        row1 = tk.Frame(self.content_frame, bg=self.colors['light'])
        row1.pack(fill=tk.X, padx=10, pady=(4, 0))
        for title, desc, cb in [
            ("📄 Word → PDF",  "Word文档转PDF文件",  self._word_to_pdf_dlg),
            ("📊 Excel → PDF", "Excel表格转PDF文件", self._excel_to_pdf_dlg),
            ("🖼 图片 → PDF",  "多张图片合成PDF",    self._images_to_pdf_dlg),
            ("📄 PDF → 图片",  "PDF每页导出为图片",  self._pdf_to_img_dlg),
        ]:
            self._create_card(row1, title, desc, cb)

        # 第二行：Excel ↔ CSV/JSON
        row2 = tk.Frame(self.content_frame, bg=self.colors['light'])
        row2.pack(fill=tk.X, padx=10, pady=4)
        for title, desc, cb in [
            ("📊 Excel → CSV",  "Excel工作簿转CSV",    self._excel_to_csv_dlg),
            ("📄 CSV → Excel",  "CSV文件转Excel表格",   self._csv_to_excel_dlg),
            ("📊 Excel → JSON", "Excel数据转JSON格式",  self._excel_to_json_dlg),
            ("📋 JSON → Excel", "JSON数据转Excel表格",  self._json_to_excel_dlg),
        ]:
            self._create_card(row2, title, desc, cb)

        # 第三行：HTML / 混合 / 批量
        row3 = tk.Frame(self.content_frame, bg=self.colors['light'])
        row3.pack(fill=tk.X, padx=10, pady=(0, 4))
        for title, desc, cb in [
            ("📊 Excel → HTML", "Excel表格导出为HTML",  self._excel_to_html_dlg),
            ("🔄 CSV ↔ JSON",   "CSV与JSON互转",        self._csv_json_convert_dlg),
            ("📄 Word → HTML",  "Word文档转HTML网页",   self._word_to_html_dlg),
            ("🖼 图片批量转",   "批量转换图片格式",     self._batch_image_convert_dlg),
        ]:
            self._create_card(row3, title, desc, cb)

    # ══════════════════════════════════════════════════════════════════════
    #  Word → PDF
    # ══════════════════════════════════════════════════════════════════════

    def _word_to_pdf_dlg(self):
        """Word → PDF：优先 Win32 COM，回退 reportlab。"""
        win = tk.Toplevel(self.root)
        win.title("Word → PDF 转换")
        win.geometry("520x300")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="📄 Word → PDF 转换", font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text="选择 Word 文档 (.docx)，转换为 PDF", font=("微软雅黑", 9), fg="gray").pack()

        # 文件选择
        file_var = tk.StringVar()
        f_frame = tk.Frame(win)
        f_frame.pack(fill=tk.X, padx=20, pady=10)
        tk.Entry(f_frame, textvariable=file_var, width=40, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(f_frame, text="浏览...", command=lambda: file_var.set(
            filedialog.askopenfilename(
                title="选择 Word 文件",
                filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
            ) or file_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        # 状态 / 日志
        log = scrolledtext.ScrolledText(win, height=6, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        def _do_convert():
            fp = file_var.get().strip()
            if not fp or not os.path.isfile(fp):
                messagebox.showwarning("提示", "请先选择有效的 Word 文件", parent=win)
                return
            out_fp = str(Path(fp).with_suffix(".pdf"))
            out_fp_save = filedialog.asksaveasfilename(
                title="保存 PDF 为...",
                initialdir=os.path.dirname(fp),
                initialfile=os.path.basename(out_fp),
                defaultextension=".pdf",
                filetypes=[("PDF文件", "*.pdf")],
                parent=win,
            )
            if not out_fp_save:
                return

            def _work():
                try:
                    log.insert(tk.END, "正在转换...\n")
                    log.see(tk.END)
                    win.update()

                    if WORD_COM_AVAILABLE:
                        self._word_to_pdf_com(fp, out_fp_save, log)
                    elif REPORTLAB_AVAILABLE:
                        self._word_to_pdf_reportlab(fp, out_fp_save, log)
                    else:
                        log.insert(tk.END, "❌ 缺少转换引擎：请安装 pywin32 (COM) 或 reportlab\n")
                        return

                    log.insert(tk.END, f"✅ 转换成功！→ {out_fp_save}\n")
                    self.set_status(f"Word→PDF: {Path(fp).name}")
                except Exception as e:
                    log.insert(tk.END, f"❌ 转换失败: {e}\n")

            threading.Thread(target=_work, daemon=True).start()

        tk.Button(win, text="🚀 开始转换", command=_do_convert, cursor="hand2",
                  bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=14).pack(pady=8)

    def _word_to_pdf_com(self, src: str, dst: str, log_widget=None):
        """使用 Win32 COM (Word.Application) 转换 docx → PDF。"""
        win32 = _import_win32com()
        word = None
        doc = None
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(os.path.abspath(src))
            # wdFormatPDF = 17
            doc.SaveAs(os.path.abspath(dst), FileFormat=17)
        except Exception as e:
            if log_widget:
                log_widget.insert(tk.END, f"⚠ COM 转换失败: {e}，尝试回退...\n")
                log_widget.see(tk.END)
            raise
        finally:
            # 清理 COM 对象（独立 try/except，不让 Close/Quit 异常掩盖成功）
            try:
                if doc:
                    doc.Close()
            except Exception:
                pass
            try:
                if word:
                    word.Quit()
            except Exception:
                pass

    def _word_to_pdf_reportlab(self, src: str, dst: str, log_widget=None):
        """使用 reportlab 重建 Word → PDF（读取 docx 文本内容后渲染）。"""
        Document = _import_docx()
        if Document is None:
            raise RuntimeError("python-docx 未安装，无法读取 Word 文件")

        rl = _import_reportlab()
        if rl is None:
            raise RuntimeError("reportlab 未安装")

        (A4, mm, canvas, HexColor, black, SimpleDocTemplate, Paragraph,
         Spacer, RLImage, Table, TableStyle, getSampleStyleSheet,
         ParagraphStyle, TA_CENTER, TA_LEFT) = rl

        register_chinese_font()
        font_name = get_chinese_font_name()

        doc = Document(src)
        pdf = SimpleDocTemplate(dst, pagesize=A4,
                                leftMargin=25*mm, rightMargin=25*mm,
                                topMargin=20*mm, bottomMargin=20*mm)

        styles = getSampleStyleSheet()
        # 注册中文字体样式
        cn_style = ParagraphStyle(
            "ChineseBody",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
        cn_heading = ParagraphStyle(
            "ChineseHeading",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceAfter=10,
        )

        elements = []
        for para in doc.paragraphs:
            text = esc_xml(para.text.strip())
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                elements.append(Paragraph(text, cn_heading))
            else:
                elements.append(Paragraph(text, cn_style))
            elements.append(Spacer(1, 4))

        # 表格
        for table in doc.tables:
            data = []
            for row in table.rows:
                data.append([esc_xml(cell.text.strip()) for cell in row.cells])
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4F46E5")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                ]))
                elements.append(Spacer(1, 8))
                elements.append(t)
                elements.append(Spacer(1, 8))

        pdf.build(elements)

    # ══════════════════════════════════════════════════════════════════════
    #  Excel → PDF
    # ══════════════════════════════════════════════════════════════════════

    def _excel_to_pdf_dlg(self):
        """Excel → PDF：openpyxl 读取 + reportlab 渲染。"""
        win = tk.Toplevel(self.root)
        win.title("Excel → PDF 转换")
        win.geometry("520x300")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="📊 Excel → PDF 转换", font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text="选择 Excel 文件 (.xlsx)，转换为 PDF", font=("微软雅黑", 9), fg="gray").pack()

        file_var = tk.StringVar()
        f_frame = tk.Frame(win)
        f_frame.pack(fill=tk.X, padx=20, pady=10)
        tk.Entry(f_frame, textvariable=file_var, width=40, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(f_frame, text="浏览...", command=lambda: file_var.set(
            filedialog.askopenfilename(
                title="选择 Excel 文件",
                filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
            ) or file_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        log = scrolledtext.ScrolledText(win, height=6, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        def _do_convert():
            fp = file_var.get().strip()
            if not fp or not os.path.isfile(fp):
                messagebox.showwarning("提示", "请先选择有效的 Excel 文件", parent=win)
                return
            out_fp_save = filedialog.asksaveasfilename(
                title="保存 PDF 为...",
                initialdir=os.path.dirname(fp),
                initialfile=Path(fp).stem + ".pdf",
                defaultextension=".pdf",
                filetypes=[("PDF文件", "*.pdf")],
                parent=win,
            )
            if not out_fp_save:
                return

            def _work():
                try:
                    log.insert(tk.END, "正在读取 Excel...\n")
                    win.update()

                    openpyxl_mod = _import_openpyxl()
                    if openpyxl_mod is None:
                        log.insert(tk.END, f"❌ 缺少 openpyxl：{_OPENPYL_ERROR}\n")
                        return
                    rl = _import_reportlab()
                    if rl is None:
                        log.insert(tk.END, f"❌ 缺少 reportlab：{_REPORTLAB_ERROR}\n")
                        return

                    (A4, mm, canvas_cls, HexColor, black, SimpleDocTemplate, Paragraph,
                     Spacer, RLImage, Table, TableStyle, getSampleStyleSheet,
                     ParagraphStyle, TA_CENTER, TA_LEFT) = rl

                    register_chinese_font()
                    font_name = get_chinese_font_name()
                    wb = openpyxl_mod.load_workbook(fp, data_only=True)
                    ws = wb.active

                    pdf = SimpleDocTemplate(out_fp_save, pagesize=A4,
                                            leftMargin=15*mm, rightMargin=15*mm,
                                            topMargin=15*mm, bottomMargin=15*mm)
                    styles = getSampleStyleSheet()
                    cell_style = ParagraphStyle(
                        "CellStyle", parent=styles["Normal"],
                        fontName=font_name, fontSize=8, leading=10,
                    )
                    header_style = ParagraphStyle(
                        "HeaderStyle", parent=styles["Normal"],
                        fontName=font_name, fontSize=9, leading=11,
                        textColor=HexColor("#FFFFFF"),
                    )

                    elements = []
                    # 标题
                    title = Paragraph(esc_xml(ws.title or "Sheet1"),
                                      ParagraphStyle("Title", parent=styles["Title"],
                                                     fontName=font_name, fontSize=14))
                    elements.append(title)
                    elements.append(Spacer(1, 10))

                    # 数据表
                    data_rows = []
                    for row in ws.iter_rows(values_only=True):
                        data_rows.append([
                            Paragraph(esc_xml(safe_str(cell)), cell_style)
                            for cell in row
                        ])

                    if data_rows:
                        t = Table(data_rows, repeatRows=1)
                        tbl_style = [
                            ("FONTNAME", (0, 0), (-1, -1), font_name),
                            ("FONTSIZE", (0, 0), (-1, -1), 8),
                            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4F46E5")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                        t.setStyle(TableStyle(tbl_style))
                        elements.append(t)

                    wb.close()
                    pdf.build(elements)
                    log.insert(tk.END, f"✅ 转换成功！→ {out_fp_save}\n")
                    self.set_status(f"Excel→PDF: {Path(fp).name}")
                except Exception as e:
                    log.insert(tk.END, f"❌ 转换失败: {e}\n")

            threading.Thread(target=_work, daemon=True).start()

        tk.Button(win, text="🚀 开始转换", command=_do_convert, cursor="hand2",
                  bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=14).pack(pady=8)

    # ══════════════════════════════════════════════════════════════════════
    #  图片 → PDF
    # ══════════════════════════════════════════════════════════════════════

    def _images_to_pdf_dlg(self):
        """多张图片合成 PDF：PIL + reportlab。"""
        win = tk.Toplevel(self.root)
        win.title("图片 → PDF 合成")
        win.geometry("560x380")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="🖼 图片 → PDF 合成", font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text="选择多张图片，合并为一个 PDF 文件", font=("微软雅黑", 9), fg="gray").pack()

        # 图片列表
        listbox = tk.Listbox(win, font=("Consolas", 9), selectmode=tk.EXTENDED)
        listbox.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        btn_frame = tk.Frame(win)
        btn_frame.pack(fill=tk.X, padx=15, pady=(0, 5))

        def _add_images():
            files = filedialog.askopenfilenames(
                title="选择图片",
                filetypes=[("图片文件", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff"), ("所有文件", "*.*")],
                parent=win,
            )
            for f in files:
                if f not in listbox.get(0, tk.END):
                    listbox.insert(tk.END, f)

        def _remove_selected():
            sel = listbox.curselection()
            for i in reversed(sel):
                listbox.delete(i)

        def _clear_list():
            listbox.delete(0, tk.END)

        def _move_up():
            sel = listbox.curselection()
            if not sel or sel[0] == 0:
                return
            idx = sel[0]
            text = listbox.get(idx)
            listbox.delete(idx)
            listbox.insert(idx - 1, text)
            listbox.selection_set(idx - 1)

        def _move_down():
            sel = listbox.curselection()
            if not sel or sel[0] >= listbox.size() - 1:
                return
            idx = sel[0]
            text = listbox.get(idx)
            listbox.delete(idx)
            listbox.insert(idx + 1, text)
            listbox.selection_set(idx + 1)

        tk.Button(btn_frame, text="➕ 添加图片", command=_add_images,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text="➖ 移除", command=_remove_selected,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text="🗑 清空", command=_clear_list,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text="↑ 上移", command=_move_up,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text="↓ 下移", command=_move_down,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)

        def _do_convert():
            if listbox.size() == 0:
                messagebox.showwarning("提示", "请先添加至少一张图片", parent=win)
                return
            out_fp = filedialog.asksaveasfilename(
                title="保存 PDF 为...",
                initialfile="output.pdf",
                defaultextension=".pdf",
                filetypes=[("PDF文件", "*.pdf")],
                parent=win,
            )
            if not out_fp:
                return

            def _work():
                try:
                    Image = _import_pil()
                    if Image is None:
                        messagebox.showerror("错误", f"Pillow 未安装: {_PIL_ERROR}", parent=win)
                        return

                    images = []
                    for i in range(listbox.size()):
                        img_path = listbox.get(i)
                        img = Image.open(img_path).convert("RGB")
                        images.append(img)

                    if images:
                        images[0].save(out_fp, save_all=True, append_images=images[1:])
                        messagebox.showinfo("成功", f"✅ PDF 已保存：{out_fp}", parent=win)
                        self.set_status(f"图片→PDF: {len(images)} 页")
                except Exception as e:
                    messagebox.showerror("错误", f"转换失败: {e}", parent=win)

            threading.Thread(target=_work, daemon=True).start()

        tk.Button(win, text="🚀 合成 PDF", command=_do_convert, cursor="hand2",
                  bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=14).pack(pady=5)

    # ══════════════════════════════════════════════════════════════════════
    #  PDF → 图片
    # ══════════════════════════════════════════════════════════════════════

    def _pdf_to_img_dlg(self):
        """PDF 每页导出为图片（pypdfium2）。"""
        win = tk.Toplevel(self.root)
        win.title("PDF → 图片 导出")
        win.geometry("520x320")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="📄 PDF → 图片 导出", font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text="选择 PDF，将每一页导出为 PNG 图片", font=("微软雅黑", 9), fg="gray").pack()

        file_var = tk.StringVar()
        f_frame = tk.Frame(win)
        f_frame.pack(fill=tk.X, padx=20, pady=10)
        tk.Entry(f_frame, textvariable=file_var, width=40, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(f_frame, text="浏览...", command=lambda: file_var.set(
            filedialog.askopenfilename(
                title="选择 PDF 文件",
                filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
            ) or file_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        # 输出目录
        out_dir_var = tk.StringVar()
        od_frame = tk.Frame(win)
        od_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(od_frame, text="输出目录:", font=("微软雅黑", 9)).pack(side=tk.LEFT)
        tk.Entry(od_frame, textvariable=out_dir_var, width=30, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=5)
        tk.Button(od_frame, text="选择...", command=lambda: out_dir_var.set(
            filedialog.askdirectory(title="选择输出目录", parent=win) or out_dir_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        # DPI 设置
        dpi_var = tk.IntVar(value=200)
        dp_frame = tk.Frame(win)
        dp_frame.pack(fill=tk.X, padx=20, pady=5)
        tk.Label(dp_frame, text="DPI:", font=("微软雅黑", 9)).pack(side=tk.LEFT)
        tk.Scale(dp_frame, from_=72, to=600, orient=tk.HORIZONTAL, variable=dpi_var,
                 length=200, font=("微软雅黑", 8)).pack(side=tk.LEFT, padx=5)
        tk.Label(dp_frame, textvariable=dpi_var, font=("微软雅黑", 9), fg="green").pack(side=tk.LEFT)

        log = scrolledtext.ScrolledText(win, height=5, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        def _do_convert():
            fp = file_var.get().strip()
            out_dir = out_dir_var.get().strip()
            if not fp or not os.path.isfile(fp):
                messagebox.showwarning("提示", "请先选择 PDF 文件", parent=win)
                return
            if not out_dir:
                out_dir = os.path.dirname(fp)
            os.makedirs(out_dir, exist_ok=True)

            def _work():
                try:
                    pdfium = _import_pypdfium2()
                    if pdfium is None:
                        log.insert(tk.END, f"❌ 缺少 pypdfium2：{_PYPDFIUM_ERROR}\n")
                        return
                    Image = _import_pil()
                    if Image is None:
                        log.insert(tk.END, f"❌ 缺少 Pillow：{_PIL_ERROR}\n")
                        return

                    dpi = dpi_var.get()
                    pdf_doc = pdfium.PdfDocument(fp)
                    total = len(pdf_doc)
                    log.insert(tk.END, f"📄 PDF 共 {total} 页，开始导出 (DPI={dpi})...\n")
                    win.update()

                    for i in range(total):
                        page = pdf_doc[i]
                        bitmap = page.render(scale=dpi / 72)
                        pil_image = bitmap.to_pil()
                        out_path = os.path.join(out_dir, f"{Path(fp).stem}_page_{i + 1:03d}.png")
                        pil_image.save(out_path, "PNG")
                        log.insert(tk.END, f"  ✅ 第 {i + 1}/{total} 页 → {Path(out_path).name}\n")
                        win.update()

                    pdf_doc.close()
                    log.insert(tk.END, f"\n🎉 导出完成！共 {total} 页\n")
                    self.set_status(f"PDF→图片: {total} 页")
                except Exception as e:
                    log.insert(tk.END, f"❌ 导出失败: {e}\n")

            threading.Thread(target=_work, daemon=True).start()

        tk.Button(win, text="🚀 开始导出", command=_do_convert, cursor="hand2",
                  bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=14).pack(pady=8)

    # ══════════════════════════════════════════════════════════════════════
    #  Excel ↔ CSV
    # ══════════════════════════════════════════════════════════════════════

    def _excel_to_csv_dlg(self):
        """Excel → CSV 转换。"""
        win = self._base_file_convert_dlg(
            title="Excel → CSV 转换",
            heading="📊 Excel → CSV 转换",
            desc="选择 Excel 文件 (.xlsx)，导出为 CSV",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")],
            ext=".csv",
            convert_fn=self._do_excel_to_csv,
        )

    def _do_excel_to_csv(self, src: str, dst: str, log_widget) -> bool:
        openpyxl_mod = _import_openpyxl()
        if openpyxl_mod is None:
            log_widget.insert(tk.END, f"❌ 缺少 openpyxl：{_OPENPYL_ERROR}\n")
            return False
        try:
            wb = openpyxl_mod.load_workbook(src, data_only=True)
            ws = wb.active
            with open(dst, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow([safe_str(c) for c in row])
            wb.close()
            log_widget.insert(tk.END, f"✅ CSV 已保存: {dst}\n")
            return True
        except Exception as e:
            log_widget.insert(tk.END, f"❌ 转换失败: {e}\n")
            return False

    def _csv_to_excel_dlg(self):
        """CSV → Excel 转换。"""
        win = self._base_file_convert_dlg(
            title="CSV → Excel 转换",
            heading="📄 CSV → Excel 转换",
            desc="选择 CSV 文件，保存为 Excel (.xlsx)",
            filetypes=[("CSV文件", "*.csv"), ("所有文件", "*.*")],
            ext=".xlsx",
            convert_fn=self._do_csv_to_excel,
        )

    def _do_csv_to_excel(self, src: str, dst: str, log_widget) -> bool:
        openpyxl_mod = _import_openpyxl()
        if openpyxl_mod is None:
            log_widget.insert(tk.END, f"❌ 缺少 openpyxl：{_OPENPYL_ERROR}\n")
            return False
        try:
            wb = openpyxl_mod.Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            with open(src, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                for row_idx, row in enumerate(reader, 1):
                    for col_idx, val in enumerate(row, 1):
                        ws.cell(row=row_idx, column=col_idx, value=val)
            wb.save(dst)
            wb.close()
            log_widget.insert(tk.END, f"✅ Excel 已保存: {dst}\n")
            return True
        except Exception as e:
            log_widget.insert(tk.END, f"❌ 转换失败: {e}\n")
            return False

    # ══════════════════════════════════════════════════════════════════════
    #  Excel ↔ JSON
    # ══════════════════════════════════════════════════════════════════════

    def _excel_to_json_dlg(self):
        """Excel → JSON 转换。"""
        win = self._base_file_convert_dlg(
            title="Excel → JSON 转换",
            heading="📊 Excel → JSON 转换",
            desc="选择 Excel 文件，导出为 JSON",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")],
            ext=".json",
            convert_fn=self._do_excel_to_json,
        )

    def _do_excel_to_json(self, src: str, dst: str, log_widget) -> bool:
        openpyxl_mod = _import_openpyxl()
        if openpyxl_mod is None:
            log_widget.insert(tk.END, f"❌ 缺少 openpyxl：{_OPENPYL_ERROR}\n")
            return False
        try:
            wb = openpyxl_mod.load_workbook(src, data_only=True)
            ws = wb.active
            headers = [safe_str(cell) for cell in next(ws.iter_rows(values_only=True))]
            data = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                item = {}
                for i, val in enumerate(row):
                    if i < len(headers):
                        item[headers[i]] = safe_str(val)
                data.append(item)
            wb.close()
            with open(dst, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            log_widget.insert(tk.END, f"✅ JSON 已保存: {dst} ({len(data)} 条记录)\n")
            return True
        except Exception as e:
            log_widget.insert(tk.END, f"❌ 转换失败: {e}\n")
            return False

    def _json_to_excel_dlg(self):
        """JSON → Excel 转换。"""
        win = self._base_file_convert_dlg(
            title="JSON → Excel 转换",
            heading="📋 JSON → Excel 转换",
            desc="选择 JSON 文件，保存为 Excel (.xlsx)",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")],
            ext=".xlsx",
            convert_fn=self._do_json_to_excel,
        )

    def _do_json_to_excel(self, src: str, dst: str, log_widget) -> bool:
        openpyxl_mod = _import_openpyxl()
        if openpyxl_mod is None:
            log_widget.insert(tk.END, f"❌ 缺少 openpyxl：{_OPENPYL_ERROR}\n")
            return False
        try:
            with open(src, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                data = [data]
            elif not isinstance(data, list):
                data = [{"value": safe_str(data)}]

            wb = openpyxl_mod.Workbook()
            ws = wb.active
            ws.title = "Sheet1"

            # 提取所有键作为表头
            all_keys = []
            for item in data:
                if isinstance(item, dict):
                    for k in item:
                        if k not in all_keys:
                            all_keys.append(k)

            if all_keys:
                for col_idx, key in enumerate(all_keys, 1):
                    ws.cell(row=1, column=col_idx, value=key)
                for row_idx, item in enumerate(data, 2):
                    if isinstance(item, dict):
                        for col_idx, key in enumerate(all_keys, 1):
                            ws.cell(row=row_idx, column=col_idx,
                                    value=safe_str(item.get(key, "")))
                    else:
                        ws.cell(row=row_idx, column=1, value=safe_str(item))
            else:
                for row_idx, item in enumerate(data, 1):
                    ws.cell(row=row_idx, column=1, value=safe_str(item))

            wb.save(dst)
            wb.close()
            log_widget.insert(tk.END, f"✅ Excel 已保存: {dst} ({len(data)} 条记录)\n")
            return True
        except Exception as e:
            log_widget.insert(tk.END, f"❌ 转换失败: {e}\n")
            return False

    # ══════════════════════════════════════════════════════════════════════
    #  Excel → HTML
    # ══════════════════════════════════════════════════════════════════════

    def _excel_to_html_dlg(self):
        """Excel → HTML 转换。"""
        win = self._base_file_convert_dlg(
            title="Excel → HTML 转换",
            heading="📊 Excel → HTML 转换",
            desc="选择 Excel 文件，导出为 HTML 表格",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")],
            ext=".html",
            convert_fn=self._do_excel_to_html,
        )

    def _do_excel_to_html(self, src: str, dst: str, log_widget) -> bool:
        openpyxl_mod = _import_openpyxl()
        if openpyxl_mod is None:
            log_widget.insert(tk.END, f"❌ 缺少 openpyxl：{_OPENPYL_ERROR}\n")
            return False
        try:
            wb = openpyxl_mod.load_workbook(src, data_only=True)
            ws = wb.active

            html_parts = [
                "<!DOCTYPE html>",
                '<html lang="zh-CN">',
                "<head><meta charset='utf-8'><title>Excel 导出</title>",
                "<style>",
                "body { font-family: 'Microsoft YaHei', sans-serif; margin: 20px; }",
                "table { border-collapse: collapse; width: 100%; }",
                "th, td { border: 1px solid #ccc; padding: 6px 10px; text-align: left; }",
                "th { background: #4F46E5; color: white; }",
                "tr:nth-child(even) { background: #F8FAFC; }",
                "</style></head><body>",
                f"<h2>{esc_xml(ws.title)}</h2>",
                "<table>",
            ]

            for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                html_parts.append("<tr>")
                tag = "th" if row_idx == 0 else "td"
                for cell in row:
                    val = esc_xml(safe_str(cell))
                    html_parts.append(f"<{tag}>{val}</{tag}>")
                html_parts.append("</tr>")

            html_parts.append("</table></body></html>")
            wb.close()

            with open(dst, "w", encoding="utf-8") as f:
                f.write("\n".join(html_parts))
            log_widget.insert(tk.END, f"✅ HTML 已保存: {dst}\n")
            return True
        except Exception as e:
            log_widget.insert(tk.END, f"❌ 转换失败: {e}\n")
            return False

    # ══════════════════════════════════════════════════════════════════════
    #  CSV ↔ JSON 互转
    # ══════════════════════════════════════════════════════════════════════

    def _csv_json_convert_dlg(self):
        """CSV ↔ JSON 双向互转对话框。"""
        win = tk.Toplevel(self.root)
        win.title("CSV ↔ JSON 互转")
        win.geometry("520x350")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="🔄 CSV ↔ JSON 互转", font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text="选择文件，一键转换 CSV ↔ JSON", font=("微软雅黑", 9), fg="gray").pack()

        file_var = tk.StringVar()
        f_frame = tk.Frame(win)
        f_frame.pack(fill=tk.X, padx=20, pady=8)
        tk.Entry(f_frame, textvariable=file_var, width=40, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(f_frame, text="浏览...", command=lambda: file_var.set(
            filedialog.askopenfilename(
                title="选择文件",
                filetypes=[("CSV/JSON", "*.csv *.json"), ("所有文件", "*.*")]
            ) or file_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        log = scrolledtext.ScrolledText(win, height=8, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        def _convert_to(mode: str):
            fp = file_var.get().strip()
            if not fp or not os.path.isfile(fp):
                messagebox.showwarning("提示", "请先选择文件", parent=win)
                return

            ext = Path(fp).suffix.lower()
            if mode == "csv2json" and ext != ".csv":
                messagebox.showwarning("提示", "请选择 CSV 文件", parent=win)
                return
            if mode == "json2csv" and ext != ".json":
                messagebox.showwarning("提示", "请选择 JSON 文件", parent=win)
                return

            out_ext = ".json" if mode == "csv2json" else ".csv"
            out_fp = filedialog.asksaveasfilename(
                title="保存文件",
                initialdir=os.path.dirname(fp),
                initialfile=Path(fp).stem + out_ext,
                defaultextension=out_ext,
                filetypes=[(out_ext.upper(), f"*{out_ext}")],
                parent=win,
            )
            if not out_fp:
                return

            def _work():
                try:
                    if mode == "csv2json":
                        data = []
                        with open(fp, "r", encoding="utf-8-sig") as f:
                            reader = csv.DictReader(f)
                            for row in reader:
                                data.append(row)
                        with open(out_fp, "w", encoding="utf-8") as f:
                            json.dump(data, f, ensure_ascii=False, indent=2)
                        log.insert(tk.END, f"✅ CSV → JSON 完成！({len(data)} 条记录)\n")
                    else:
                        with open(fp, "r", encoding="utf-8") as f:
                            data = json.load(f)
                        if isinstance(data, dict):
                            data = [data]
                        with open(out_fp, "w", newline="", encoding="utf-8-sig") as f:
                            if data and isinstance(data[0], dict):
                                writer = csv.DictWriter(f, fieldnames=list(data[0].keys()))
                                writer.writeheader()
                                writer.writerows(data)
                            else:
                                writer = csv.writer(f)
                                for item in data:
                                    writer.writerow([safe_str(item)])
                        log.insert(tk.END, f"✅ JSON → CSV 完成！({len(data)} 条记录)\n")
                    log.insert(tk.END, f"   → {out_fp}\n")
                    self.set_status(f"CSV↔JSON: {Path(fp).name}")
                except Exception as e:
                    log.insert(tk.END, f"❌ 转换失败: {e}\n")

            threading.Thread(target=_work, daemon=True).start()

        btn_f = tk.Frame(win)
        btn_f.pack(pady=5)
        tk.Button(btn_f, text="CSV → JSON", command=lambda: _convert_to("csv2json"),
                  cursor="hand2", bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=12).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_f, text="JSON → CSV", command=lambda: _convert_to("json2csv"),
                  cursor="hand2", bg=self.colors['secondary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=12).pack(side=tk.LEFT, padx=5)

    # ══════════════════════════════════════════════════════════════════════
    #  Word → HTML
    # ══════════════════════════════════════════════════════════════════════

    def _word_to_html_dlg(self):
        """Word → HTML 转换。"""
        win = self._base_file_convert_dlg(
            title="Word → HTML 转换",
            heading="📄 Word → HTML 转换",
            desc="选择 Word 文档 (.docx)，导出为 HTML",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            ext=".html",
            convert_fn=self._do_word_to_html,
        )

    def _do_word_to_html(self, src: str, dst: str, log_widget) -> bool:
        Document = _import_docx()
        if Document is None:
            log_widget.insert(tk.END, f"❌ 缺少 python-docx：{_DOCX_ERROR}\n")
            return False
        try:
            doc = Document(src)
            html_parts = [
                "<!DOCTYPE html>",
                '<html lang="zh-CN">',
                "<head><meta charset='utf-8'><title>Word 导出</title>",
                "<style>",
                "body { font-family: 'Microsoft YaHei', sans-serif; margin: 30px; line-height: 1.6; }",
                "h1, h2, h3 { color: #1E293B; }",
                "table { border-collapse: collapse; width: 100%; margin: 10px 0; }",
                "th, td { border: 1px solid #ccc; padding: 6px 10px; }",
                "th { background: #4F46E5; color: white; }",
                "p { margin: 6px 0; }",
                "</style></head><body>",
            ]

            for para in doc.paragraphs:
                text = esc_xml(para.text.strip())
                if not text:
                    continue
                style_name = para.style.name.lower() if para.style else ""
                if "heading 1" in style_name:
                    html_parts.append(f"<h1>{text}</h1>")
                elif "heading 2" in style_name:
                    html_parts.append(f"<h2>{text}</h2>")
                elif "heading 3" in style_name:
                    html_parts.append(f"<h3>{text}</h3>")
                else:
                    html_parts.append(f"<p>{text}</p>")

            for table in doc.tables:
                html_parts.append("<table>")
                for row_idx, row in enumerate(table.rows):
                    html_parts.append("<tr>")
                    tag = "th" if row_idx == 0 else "td"
                    for cell in row.cells:
                        html_parts.append(f"<{tag}>{esc_xml(cell.text.strip())}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

            html_parts.append("</body></html>")
            with open(dst, "w", encoding="utf-8") as f:
                f.write("\n".join(html_parts))
            log_widget.insert(tk.END, f"✅ HTML 已保存: {dst}\n")
            return True
        except Exception as e:
            log_widget.insert(tk.END, f"❌ 转换失败: {e}\n")
            return False

    # ══════════════════════════════════════════════════════════════════════
    #  图片批量格式转换
    # ══════════════════════════════════════════════════════════════════════

    def _batch_image_convert_dlg(self):
        """批量转换图片格式。"""
        win = tk.Toplevel(self.root)
        win.title("批量图片格式转换")
        win.geometry("560x420")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="🖼 批量图片格式转换", font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text="批量转换图片为指定格式（支持 PNG / JPG / BMP / GIF / WEBP）",
                 font=("微软雅黑", 9), fg="gray").pack()

        listbox = tk.Listbox(win, font=("Consolas", 9), selectmode=tk.EXTENDED)
        listbox.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        btn_f = tk.Frame(win)
        btn_f.pack(fill=tk.X, padx=15)

        def _add():
            files = filedialog.askopenfilenames(
                title="选择图片",
                filetypes=[("图片文件", "*.png *.jpg *.jpeg *.bmp *.gif *.webp *.tiff"), ("所有文件", "*.*")],
                parent=win,
            )
            for f in files:
                if f not in listbox.get(0, tk.END):
                    listbox.insert(tk.END, f)

        def _remove():
            sel = listbox.curselection()
            for i in reversed(sel):
                listbox.delete(i)

        def _clear():
            listbox.delete(0, tk.END)

        tk.Button(btn_f, text="➕ 添加图片", command=_add,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_f, text="➖ 移除", command=_remove,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_f, text="🗑 清空", command=_clear,
                  cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)

        # 目标格式选择
        fmt_var = tk.StringVar(value="PNG")
        fmt_frame = tk.Frame(win)
        fmt_frame.pack(fill=tk.X, padx=15, pady=5)
        tk.Label(fmt_frame, text="目标格式:", font=("微软雅黑", 9)).pack(side=tk.LEFT)
        fmt_combo = ttk.Combobox(fmt_frame, textvariable=fmt_var,
                                 values=["PNG", "JPG", "BMP", "GIF", "WEBP"],
                                 state="readonly", width=8, font=("微软雅黑", 9))
        fmt_combo.pack(side=tk.LEFT, padx=5)

        # 输出目录
        out_dir_var = tk.StringVar()
        tk.Label(fmt_frame, text="输出目录:", font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=(10, 0))
        tk.Entry(fmt_frame, textvariable=out_dir_var, width=18, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=3)
        tk.Button(fmt_frame, text="选择...", command=lambda: out_dir_var.set(
            filedialog.askdirectory(title="选择输出目录", parent=win) or out_dir_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        log = scrolledtext.ScrolledText(win, height=5, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        def _do_convert():
            if listbox.size() == 0:
                messagebox.showwarning("提示", "请先添加图片", parent=win)
                return
            out_dir = out_dir_var.get().strip()
            if not out_dir:
                out_dir = os.path.dirname(listbox.get(0))
            os.makedirs(out_dir, exist_ok=True)
            target_fmt = fmt_var.get().lower()

            def _work():
                try:
                    Image = _import_pil()
                    if Image is None:
                        log.insert(tk.END, f"❌ 缺少 Pillow: {_PIL_ERROR}\n")
                        return

                    total = listbox.size()
                    success = 0
                    for i in range(total):
                        src_path = listbox.get(i)
                        try:
                            img = Image.open(src_path)
                            stem = Path(src_path).stem
                            out_path = os.path.join(out_dir, f"{stem}.{target_fmt}")
                            # 特殊处理：JPG 不支持 RGBA
                            if target_fmt in ("jpg", "jpeg") and img.mode in ("RGBA", "P"):
                                img = img.convert("RGB")
                            img.save(out_path, format=target_fmt.upper())
                            success += 1
                            log.insert(tk.END, f"  ✅ ({i + 1}/{total}) {stem}.{target_fmt}\n")
                        except Exception as e:
                            log.insert(tk.END, f"  ❌ ({i + 1}/{total}) {Path(src_path).name}: {e}\n")
                        win.update()

                    log.insert(tk.END, f"\n🎉 完成！成功 {success}/{total} 张\n")
                    self.set_status(f"图片批量转{target_fmt}: {success}/{total}")
                except Exception as e:
                    log.insert(tk.END, f"❌ 批量转换失败: {e}\n")

            threading.Thread(target=_work, daemon=True).start()

        tk.Button(win, text="🚀 开始转换", command=_do_convert, cursor="hand2",
                  bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=14).pack(pady=8)

    # ══════════════════════════════════════════════════════════════════════
    #  通用：文件转换基础对话框
    # ══════════════════════════════════════════════════════════════════════

    def _base_file_convert_dlg(self, *, title: str, heading: str, desc: str,
                                filetypes: list, ext: str,
                                convert_fn) -> tk.Toplevel:
        """创建标准的「选择文件 → 转换 → 保存」对话框。

        参数:
            title:     窗口标题
            heading:   顶部大标题（可含 emoji）
            desc:      描述文字
            filetypes: filedialog 的文件类型过滤
            ext:       目标扩展名（含点，如 .csv）
            convert_fn: 转换函数，签名 convert_fn(src, dst, log_widget) -> bool

        返回:
            tk.Toplevel 实例
        """
        win = tk.Toplevel(self.root)
        win.title(title)
        win.geometry("520x280")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text=heading, font=("微软雅黑", 14, "bold")).pack(pady=8)
        tk.Label(win, text=desc, font=("微软雅黑", 9), fg="gray").pack()

        file_var = tk.StringVar()
        f_frame = tk.Frame(win)
        f_frame.pack(fill=tk.X, padx=20, pady=10)
        tk.Entry(f_frame, textvariable=file_var, width=40, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(f_frame, text="浏览...", command=lambda: file_var.set(
            filedialog.askopenfilename(title=title, filetypes=filetypes, parent=win) or file_var.get()
        ), cursor="hand2", font=("微软雅黑", 9)).pack(side=tk.LEFT)

        log = scrolledtext.ScrolledText(win, height=5, font=("Consolas", 9))
        log.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        def _do():
            fp = file_var.get().strip()
            if not fp or not os.path.isfile(fp):
                messagebox.showwarning("提示", "请先选择有效文件", parent=win)
                return
            out_fp = filedialog.asksaveasfilename(
                title="保存文件",
                initialdir=os.path.dirname(fp),
                initialfile=Path(fp).stem + ext,
                defaultextension=ext,
                filetypes=[(ext.upper(), f"*{ext}")],
                parent=win,
            )
            if not out_fp:
                return

            def _work():
                try:
                    log.insert(tk.END, "正在转换...\n")
                    win.update()
                    ok = convert_fn(fp, out_fp, log)
                    if ok:
                        log.insert(tk.END, f"✅ 转换完成！\n")
                        self.set_status(f"{title}: {Path(fp).name}")
                except Exception as e:
                    log.insert(tk.END, f"❌ 转换失败: {e}\n")

            threading.Thread(target=_work, daemon=True).start()

        tk.Button(win, text="🚀 开始转换", command=_do, cursor="hand2",
                  bg=self.colors['primary'], fg="white",
                  font=("微软雅黑", 10, "bold"), width=14).pack(pady=5)

        return win
